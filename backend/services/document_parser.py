"""
Multi-Format Document Parser
Unterstützt: PDF, DOCX, DOC, TXT, RTF, XLS, XLSX, CSV, JPG, PNG, TIFF, WEBP, EML, XML, HTML
Mit OCR-Unterstützung für gescannte PDFs
"""

import io
import logging
import csv
from typing import Optional, Dict
from enum import Enum
from pydantic import BaseModel, Field

import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes
from docx import Document

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    """Supported document types"""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    IMAGE = "image"  # JPG, PNG, TIFF, WEBP
    EXCEL = "excel"  # XLS, XLSX
    CSV = "csv"
    RTF = "rtf"
    HTML = "html"
    XML = "xml"
    EML = "eml"


class DocumentParseResult(BaseModel):
    """Result of document parsing"""
    text: str = Field(..., description="Extracted text content")
    doc_type: DocumentType
    page_count: Optional[int] = None
    char_count: int
    word_count: int
    ocr_applied: bool = False
    error: Optional[str] = None


class DocumentParser:
    """
    Unified document parser for all supported formats.
    
    Supported formats:
    - PDF: PyMuPDF (fitz)
    - DOCX: python-docx
    - TXT: direct read
    - Images (JPG, PNG): pytesseract OCR
    """
    
    @staticmethod
    def detect_document_type(filename: str) -> DocumentType:
        """
        Detect document type from filename extension.
        
        Args:
            filename: Original filename
            
        Returns:
            DocumentType enum
        """
        ext = filename.lower().split('.')[-1]
        
        if ext == 'pdf':
            return DocumentType.PDF
        elif ext in ['docx', 'doc']:
            return DocumentType.DOCX
        elif ext == 'txt':
            return DocumentType.TXT
        elif ext in ['jpg', 'jpeg', 'png', 'tiff', 'tif', 'webp']:
            return DocumentType.IMAGE
        elif ext in ['xlsx', 'xls']:
            return DocumentType.EXCEL
        elif ext == 'csv':
            return DocumentType.CSV
        elif ext == 'rtf':
            return DocumentType.RTF
        elif ext in ['html', 'htm']:
            return DocumentType.HTML
        elif ext == 'xml':
            return DocumentType.XML
        elif ext == 'eml':
            return DocumentType.EML
        else:
            # Fallback: Versuche als Text zu lesen
            logger.warning(f"Unknown file type: {ext}, treating as TXT")
            return DocumentType.TXT
    
    @staticmethod
    async def extract_text_with_ocr(file_bytes: bytes) -> str:
        """
        Extract text from scanned PDF using OCR (Tesseract).
        
        Args:
            file_bytes: Raw PDF bytes
            
        Returns:
            Extracted text as string
        """
        logger.info("🔍 Using OCR to extract text from scanned PDF...")
        
        try:
            # Convert PDF pages to images
            images = convert_from_bytes(file_bytes, dpi=300)
            logger.info(f"📄 Converted PDF to {len(images)} images")
            
            full_text = []
            
            for i, image in enumerate(images):
                # Run OCR on each page (German + English)
                text = pytesseract.image_to_string(image, lang='deu+eng')
                if text.strip():
                    full_text.append(text)
                logger.debug(f"📝 OCR page {i+1}: {len(text)} characters")
            
            extracted = "\n\n".join(full_text).strip()
            logger.info(f"✅ OCR extracted {len(extracted)} characters from {len(images)} pages")
            
            return extracted
            
        except Exception as e:
            logger.error(f"❌ OCR extraction failed: {e}")
            raise ValueError(f"OCR extraction failed: {str(e)}")
    
    @staticmethod
    async def parse_pdf(file_bytes: bytes) -> DocumentParseResult:
        """
        Extract text from PDF using PyMuPDF.
        Falls back to OCR for scanned documents.
        
        Args:
            file_bytes: Raw PDF bytes
            
        Returns:
            DocumentParseResult with extracted text
        """
        try:
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
            
            full_text = []
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text = page.get_text("text")
                full_text.append(text)
            
            page_count = pdf_document.page_count
            pdf_document.close()
            
            extracted = "\n\n".join(full_text).strip()
            
            # If no text or very little text, try OCR
            if not extracted or len(extracted) < 100:
                logger.info("📄 PDF has little/no extractable text, trying OCR...")
                try:
                    extracted = await DocumentParser.extract_text_with_ocr(file_bytes)
                except Exception as ocr_error:
                    logger.error(f"OCR also failed: {ocr_error}")
                    return DocumentParseResult(
                        text="",
                        doc_type=DocumentType.PDF,
                        page_count=page_count,
                        char_count=0,
                        word_count=0,
                        error=f"Konnte keinen Text extrahieren (auch nicht mit OCR): {str(ocr_error)}"
                    )
            
            if not extracted:
                return DocumentParseResult(
                    text="",
                    doc_type=DocumentType.PDF,
                    page_count=page_count,
                    char_count=0,
                    word_count=0,
                    error="Konnte keinen Text aus dem PDF extrahieren"
                )
            
            logger.info(f"✅ PDF: Extracted {len(extracted)} chars from {page_count} pages")
            
            return DocumentParseResult(
                text=extracted,
                doc_type=DocumentType.PDF,
                page_count=page_count,
                char_count=len(extracted),
                word_count=len(extracted.split())
            )
        
        except Exception as e:
            logger.error(f"❌ PDF extraction failed: {e}")
            return DocumentParseResult(
                text="",
                doc_type=DocumentType.PDF,
                char_count=0,
                word_count=0,
                error=str(e)
            )
    
    @staticmethod
    async def parse_docx(file_bytes: bytes) -> DocumentParseResult:
        """
        Extract text from DOCX using python-docx.
        
        Args:
            file_bytes: Raw DOCX bytes
            
        Returns:
            DocumentParseResult with extracted text
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)
            
            # Combine
            full_text = '\n\n'.join(paragraphs)
            if table_texts:
                full_text += '\n\nTABELLEN:\n' + '\n'.join(table_texts)
            
            extracted = full_text.strip()
            
            logger.info(f"✅ DOCX: Extracted {len(extracted)} chars from {len(doc.paragraphs)} paragraphs")
            
            return DocumentParseResult(
                text=extracted,
                doc_type=DocumentType.DOCX,
                char_count=len(extracted),
                word_count=len(extracted.split())
            )
        
        except Exception as e:
            logger.error(f"❌ DOCX extraction failed: {e}")
            return DocumentParseResult(
                text="",
                doc_type=DocumentType.DOCX,
                char_count=0,
                word_count=0,
                error=str(e)
            )
    
    @staticmethod
    async def parse_txt(file_bytes: bytes) -> DocumentParseResult:
        """
        Extract text from TXT file.
        
        Args:
            file_bytes: Raw TXT bytes
            
        Returns:
            DocumentParseResult with text
        """
        try:
            # Try UTF-8 first, then fallback to Latin-1
            try:
                text = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                text = file_bytes.decode('latin-1')
            
            extracted = text.strip()
            
            logger.info(f"✅ TXT: Extracted {len(extracted)} chars")
            
            return DocumentParseResult(
                text=extracted,
                doc_type=DocumentType.TXT,
                char_count=len(extracted),
                word_count=len(extracted.split())
            )
        
        except Exception as e:
            logger.error(f"❌ TXT extraction failed: {e}")
            return DocumentParseResult(
                text="",
                doc_type=DocumentType.TXT,
                char_count=0,
                word_count=0,
                error=str(e)
            )
    
    @staticmethod
    async def parse_image_ocr(file_bytes: bytes) -> DocumentParseResult:
        """
        Extract text from image using Tesseract OCR.
        
        Args:
            file_bytes: Raw image bytes (JPG, PNG)
            
        Returns:
            DocumentParseResult with OCR-extracted text
        """
        try:
            # Open image
            image = Image.open(io.BytesIO(file_bytes))
            
            # Convert to RGB if needed
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # OCR with German + English language
            # Config: --oem 3 = Default OCR Engine Mode, --psm 3 = Fully automatic page segmentation
            custom_config = r'--oem 3 --psm 3 -l deu+eng'
            text = pytesseract.image_to_string(image, config=custom_config)
            
            extracted = text.strip()
            
            if not extracted:
                logger.warning("OCR extracted no text from image")
                return DocumentParseResult(
                    text="",
                    doc_type=DocumentType.IMAGE,
                    char_count=0,
                    word_count=0,
                    ocr_applied=True,
                    error="No text detected in image"
                )
            
            logger.info(f"✅ OCR: Extracted {len(extracted)} chars from image")
            
            return DocumentParseResult(
                text=extracted,
                doc_type=DocumentType.IMAGE,
                char_count=len(extracted),
                word_count=len(extracted.split()),
                ocr_applied=True
            )
        
        except Exception as e:
            logger.error(f"❌ OCR extraction failed: {e}")
            return DocumentParseResult(
                text="",
                doc_type=DocumentType.IMAGE,
                char_count=0,
                word_count=0,
                ocr_applied=True,
                error=str(e)
            )
    
    @staticmethod
    async def parse_excel(file_bytes: bytes) -> DocumentParseResult:
        """
        Extract text from Excel files (XLS, XLSX) using openpyxl.
        """
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            
            all_text = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                all_text.append(f"=== Blatt: {sheet_name} ===")
                
                for row in sheet.iter_rows():
                    row_values = [str(cell.value) if cell.value is not None else '' for cell in row]
                    row_text = ' | '.join(row_values)
                    if row_text.strip(' |'):
                        all_text.append(row_text)
            
            wb.close()
            extracted = '\n'.join(all_text).strip()
            
            logger.info(f"✅ Excel: Extracted {len(extracted)} chars from {len(wb.sheetnames)} sheets")
            
            return DocumentParseResult(
                text=extracted,
                doc_type=DocumentType.EXCEL,
                char_count=len(extracted),
                word_count=len(extracted.split())
            )
        except Exception as e:
            logger.error(f"❌ Excel extraction failed: {e}")
            return DocumentParseResult(
                text="",
                doc_type=DocumentType.EXCEL,
                char_count=0,
                word_count=0,
                error=str(e)
            )
    
    @staticmethod
    async def parse_csv(file_bytes: bytes) -> DocumentParseResult:
        """
        Extract text from CSV files.
        """
        try:
            # Try UTF-8 first, then fallback
            try:
                text = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                text = file_bytes.decode('latin-1')
            
            # Parse CSV and format nicely
            reader = csv.reader(io.StringIO(text))
            all_rows = []
            for row in reader:
                all_rows.append(' | '.join(row))
            
            extracted = '\n'.join(all_rows).strip()
            
            logger.info(f"✅ CSV: Extracted {len(extracted)} chars from {len(all_rows)} rows")
            
            return DocumentParseResult(
                text=extracted,
                doc_type=DocumentType.CSV,
                char_count=len(extracted),
                word_count=len(extracted.split())
            )
        except Exception as e:
            logger.error(f"❌ CSV extraction failed: {e}")
            return DocumentParseResult(
                text="",
                doc_type=DocumentType.CSV,
                char_count=0,
                word_count=0,
                error=str(e)
            )
    
    @staticmethod
    async def parse_html(file_bytes: bytes) -> DocumentParseResult:
        """
        Extract text from HTML files using BeautifulSoup.
        """
        try:
            from bs4 import BeautifulSoup
            
            try:
                html_text = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                html_text = file_bytes.decode('latin-1')
            
            soup = BeautifulSoup(html_text, 'html.parser')
            
            # Remove script and style elements
            for element in soup(['script', 'style', 'head', 'meta', 'link']):
                element.decompose()
            
            extracted = soup.get_text(separator='\n', strip=True)
            
            logger.info(f"✅ HTML: Extracted {len(extracted)} chars")
            
            return DocumentParseResult(
                text=extracted,
                doc_type=DocumentType.HTML,
                char_count=len(extracted),
                word_count=len(extracted.split())
            )
        except Exception as e:
            logger.error(f"❌ HTML extraction failed: {e}")
            return DocumentParseResult(
                text="",
                doc_type=DocumentType.HTML,
                char_count=0,
                word_count=0,
                error=str(e)
            )
    
    @staticmethod
    async def parse_xml(file_bytes: bytes) -> DocumentParseResult:
        """
        Extract text content from XML files.
        """
        try:
            try:
                xml_text = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                xml_text = file_bytes.decode('latin-1')
            
            # Simple text extraction from XML
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(xml_text, 'xml')
            extracted = soup.get_text(separator='\n', strip=True)
            
            logger.info(f"✅ XML: Extracted {len(extracted)} chars")
            
            return DocumentParseResult(
                text=extracted,
                doc_type=DocumentType.XML,
                char_count=len(extracted),
                word_count=len(extracted.split())
            )
        except Exception as e:
            logger.error(f"❌ XML extraction failed: {e}")
            return DocumentParseResult(
                text="",
                doc_type=DocumentType.XML,
                char_count=0,
                word_count=0,
                error=str(e)
            )
    
    @staticmethod
    async def parse_eml(file_bytes: bytes) -> DocumentParseResult:
        """
        Extract text from EML (email) files.
        """
        try:
            import email
            from email import policy
            
            msg = email.message_from_bytes(file_bytes, policy=policy.default)
            
            # Extract headers
            headers = []
            if msg['Subject']:
                headers.append(f"Betreff: {msg['Subject']}")
            if msg['From']:
                headers.append(f"Von: {msg['From']}")
            if msg['To']:
                headers.append(f"An: {msg['To']}")
            if msg['Date']:
                headers.append(f"Datum: {msg['Date']}")
            
            # Extract body
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body = part.get_content()
                        break
                    elif part.get_content_type() == "text/html":
                        from bs4 import BeautifulSoup
                        html = part.get_content()
                        soup = BeautifulSoup(html, 'html.parser')
                        body = soup.get_text(separator='\n', strip=True)
            else:
                body = msg.get_content()
            
            extracted = '\n'.join(headers) + '\n\n' + str(body)
            
            logger.info(f"✅ EML: Extracted {len(extracted)} chars")
            
            return DocumentParseResult(
                text=extracted,
                doc_type=DocumentType.EML,
                char_count=len(extracted),
                word_count=len(extracted.split())
            )
        except Exception as e:
            logger.error(f"❌ EML extraction failed: {e}")
            return DocumentParseResult(
                text="",
                doc_type=DocumentType.EML,
                char_count=0,
                word_count=0,
                error=str(e)
            )
    
    @staticmethod
    async def parse_rtf(file_bytes: bytes) -> DocumentParseResult:
        """
        Extract text from RTF files.
        """
        try:
            # Simple RTF text extraction
            try:
                rtf_text = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                rtf_text = file_bytes.decode('latin-1')
            
            # Basic RTF stripping (simple approach)
            import re
            # Remove RTF control words
            text = re.sub(r'\\[a-z]+\d*\s?', '', rtf_text)
            text = re.sub(r'[{}]', '', text)
            text = re.sub(r'\s+', ' ', text)
            extracted = text.strip()
            
            logger.info(f"✅ RTF: Extracted {len(extracted)} chars")
            
            return DocumentParseResult(
                text=extracted,
                doc_type=DocumentType.RTF,
                char_count=len(extracted),
                word_count=len(extracted.split())
            )
        except Exception as e:
            logger.error(f"❌ RTF extraction failed: {e}")
            return DocumentParseResult(
                text="",
                doc_type=DocumentType.RTF,
                char_count=0,
                word_count=0,
                error=str(e)
            )
    
    async def parse_document(
        self, 
        file_bytes: bytes,
        filename: str
    ) -> DocumentParseResult:
        """
        Route document to correct parser based on file type.
        
        Args:
            file_bytes: Raw file bytes
            filename: Original filename (for type detection)
            
        Returns:
            DocumentParseResult with extracted text
        """
        try:
            doc_type = self.detect_document_type(filename)
            
            logger.info(f"📄 Parsing {doc_type.value} document: {filename}")
            
            if doc_type == DocumentType.PDF:
                return await self.parse_pdf(file_bytes)
            elif doc_type == DocumentType.DOCX:
                return await self.parse_docx(file_bytes)
            elif doc_type == DocumentType.TXT:
                return await self.parse_txt(file_bytes)
            elif doc_type == DocumentType.IMAGE:
                return await self.parse_image_ocr(file_bytes)
            elif doc_type == DocumentType.EXCEL:
                return await self.parse_excel(file_bytes)
            elif doc_type == DocumentType.CSV:
                return await self.parse_csv(file_bytes)
            elif doc_type == DocumentType.HTML:
                return await self.parse_html(file_bytes)
            elif doc_type == DocumentType.XML:
                return await self.parse_xml(file_bytes)
            elif doc_type == DocumentType.EML:
                return await self.parse_eml(file_bytes)
            elif doc_type == DocumentType.RTF:
                return await self.parse_rtf(file_bytes)
            else:
                # Fallback: treat as text
                return await self.parse_txt(file_bytes)
        
        except Exception as e:
            logger.error(f"❌ Document parsing failed: {e}")
            return DocumentParseResult(
                text="",
                doc_type=DocumentType.TXT,  # fallback
                char_count=0,
                word_count=0,
                error=str(e)
            )


# Singleton instance
document_parser = DocumentParser()
