"""
Document Export Service
Exportiert generierte Schriftsätze als DOCX oder PDF
"""

import io
import logging
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from pydantic import BaseModel

logger = logging.getLogger(__name__)


class ExportFormat(str):
    """Supported export formats"""
    DOCX = "docx"
    PDF = "pdf"


class DocumentExportRequest(BaseModel):
    """Request für Dokument-Export"""
    content: str
    template_id: str
    format: str  # "docx" oder "pdf"
    kanzlei_name: Optional[str] = None
    kanzlei_adresse: Optional[str] = None


class DocumentExporter:
    """
    Export Service für juristische Dokumente.
    Unterstützt DOCX und PDF mit professionellem Layout.
    """
    
    @staticmethod
    def export_docx(
        content: str,
        template_id: str,
        kanzlei_name: Optional[str] = None,
        kanzlei_adresse: Optional[str] = None
    ) -> bytes:
        """
        Exportiert Dokument als DOCX mit professionellem Layout.
        
        Args:
            content: Generierter Dokumententext
            template_id: ID der Vorlage (für Metadaten)
            kanzlei_name: Optional: Kanzleiname für Header
            kanzlei_adresse: Optional: Kanzleiadresse für Header
            
        Returns:
            DOCX file bytes
        """
        try:
            # Erstelle neues Document
            doc = Document()
            
            # === HEADER ===
            if kanzlei_name:
                section = doc.sections[0]
                header = section.header
                header_para = header.paragraphs[0]
                header_para.text = f"{kanzlei_name}\n{kanzlei_adresse or ''}"
                header_para.style.font.size = Pt(8)
                header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # === MAIN CONTENT ===
            
            # Parse content (split by lines)
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    # Leere Zeile = Absatz
                    doc.add_paragraph()
                    continue
                
                # Erkennung von Überschriften (ALL CAPS)
                if line.isupper() and len(line) < 100:
                    para = doc.add_paragraph(line)
                    para.style = 'Heading 1'
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.runs[0].bold = True
                    para.runs[0].font.size = Pt(14)
                else:
                    # Normaler Text
                    para = doc.add_paragraph(line)
                    para.style.font.name = 'Times New Roman'
                    para.style.font.size = Pt(12)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # === FOOTER ===
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Erstellt mit domulex.ai | {datetime.now().strftime('%d.%m.%Y')}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.style.font.size = Pt(8)
            footer_para.style.font.italic = True
            
            # === EXPORT TO BYTES ===
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            logger.info(f"✅ DOCX export successful: {template_id}")
            
            return buffer.getvalue()
        
        except Exception as e:
            logger.error(f"DOCX export failed: {e}", exc_info=True)
            raise
    
    @staticmethod
    def export_pdf(
        content: str,
        template_id: str,
        kanzlei_name: Optional[str] = None,
        kanzlei_adresse: Optional[str] = None
    ) -> bytes:
        """
        Exportiert Dokument als PDF mit professionellem Layout.
        
        Args:
            content: Generierter Dokumententext
            template_id: ID der Vorlage
            kanzlei_name: Optional: Kanzleiname für Header
            kanzlei_adresse: Optional: Kanzleiadresse für Header
            
        Returns:
            PDF file bytes
        """
        try:
            buffer = io.BytesIO()
            
            # PDF Setup
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=2.5*cm,
                leftMargin=2.5*cm,
                topMargin=3*cm,
                bottomMargin=2.5*cm
            )
            
            # Styles
            styles = getSampleStyleSheet()
            
            # Custom Styles
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=12,
                leading=16,
                fontName='Times-Roman',
                alignment=TA_LEFT
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading1'],
                fontSize=14,
                leading=18,
                fontName='Times-Bold',
                alignment=TA_CENTER,
                spaceAfter=20
            )
            
            header_style = ParagraphStyle(
                'Header',
                parent=styles['Normal'],
                fontSize=8,
                fontName='Times-Roman',
                alignment=TA_RIGHT
            )
            
            # Story (Content)
            story = []
            
            # Header
            if kanzlei_name:
                story.append(Paragraph(kanzlei_name, header_style))
                if kanzlei_adresse:
                    story.append(Paragraph(kanzlei_adresse, header_style))
                story.append(Spacer(1, 0.5*cm))
            
            # Main Content
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    story.append(Spacer(1, 0.3*cm))
                    continue
                
                # Überschriften (ALL CAPS)
                if line.isupper() and len(line) < 100:
                    story.append(Paragraph(line, heading_style))
                else:
                    # Normaler Text
                    # Escape HTML characters
                    line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(line, normal_style))
            
            # Footer
            story.append(Spacer(1, 1*cm))
            footer_text = f"Erstellt mit domulex.ai | {datetime.now().strftime('%d.%m.%Y')}"
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontSize=8,
                fontName='Times-Italic',
                alignment=TA_CENTER
            )
            story.append(Paragraph(footer_text, footer_style))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            
            logger.info(f"✅ PDF export successful: {template_id}")
            
            return buffer.getvalue()
        
        except Exception as e:
            logger.error(f"PDF export failed: {e}", exc_info=True)
            raise


# Singleton instance
document_exporter = DocumentExporter()
