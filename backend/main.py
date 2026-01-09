"""
DOMULEX Backend - FastAPI Application
"""

import time
import logging
from contextlib import asynccontextmanager
from typing import List, Optional

from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, Form, Request, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from qdrant_client import QdrantClient
import google.generativeai as genai

from models import (
    QueryRequest,
    QueryResponse,
    IngestionRequest,
    IngestionResponse,
    Jurisdiction,
    UserRole,
    ConflictRequest,
    ConflictResponse,
    ConflictAnalysis,
)
from rag import RAGEngine
from rag.personas import get_mediator_prompt
from ingestion import ScraperFactory
from config import get_settings
from services.pdf_parser import PDFParser, ContractAnalysis, ClauseAnalysis, RiskLevel
from logger import setup_logging, RequestLogger
from auth import initialize_firebase, get_current_user, register_session, FirebaseUser

# Setup logging - MUST BE BEFORE OTHER IMPORTS THAT USE LOGGER
setup_logging()
logger = logging.getLogger(__name__)

# Import Stripe service (optional, may not be available)
try:
    from services.stripe_service import StripeService
    STRIPE_ENABLED = True
except ImportError:
    STRIPE_ENABLED = False

# Import Email service
try:
    from services.email_service import email_service, EmailService
    EMAIL_ENABLED = True
except ImportError:
    EMAIL_ENABLED = False
    logger.warning("Email service not available")

# Import Email Client (for full send/receive)
try:
    from services.email_client import email_client, EmailClient
    EMAIL_CLIENT_ENABLED = True
except ImportError:
    EMAIL_CLIENT_ENABLED = False
    logger.warning("Email client not available")

# Import Document Generator service
try:
    from services.document_generator import DocumentGenerator, DocumentRequest, GeneratedDocument
    DOCUMENT_GEN_ENABLED = True
except ImportError:
    DOCUMENT_GEN_ENABLED = False

# Import CRM models (always available)
try:
    from models.crm import (
        CreateClientRequest, UpdateClientRequest, Client,
        CreateMandateRequest, UpdateMandateRequest, Mandate, AddDeadlineRequest,
        UploadDocumentRequest, SearchDocumentsRequest, DocumentSearchResult, Document,
        GenerateInsightsRequest, MandateInsights,
        ChatWithDocumentsRequest, ChatWithDocumentsResponse,
        ClientStatus, MandateStatus, DocumentCategory
    )
    CRM_MODELS_AVAILABLE = True
except ImportError as e:
    CRM_MODELS_AVAILABLE = False
    # Define placeholders to prevent NameError at module load
    class Client: pass
    class Document: pass
    class Mandate: pass
    class DocumentSearchResult: pass
    class MandateInsights: pass
    class ChatWithDocumentsResponse: pass
    class CreateClientRequest: pass
    class UpdateClientRequest: pass
    class CreateMandateRequest: pass
    class UpdateMandateRequest: pass
    class UploadDocumentRequest: pass
    class SearchDocumentsRequest: pass
    class AddDeadlineRequest: pass
    class GenerateInsightsRequest: pass
    class ChatWithDocumentsRequest: pass
    class ClientStatus: pass
    class MandateStatus: pass
    class DocumentCategory: pass
    logger.warning(f"CRM models not available: {e}")

# CRM Service - requires Firebase (initialized lazily)
CRM_ENABLED = False
CRMService = None

if not STRIPE_ENABLED:
    logger.warning("Stripe service not available - payment endpoints will return 503")
if not EMAIL_ENABLED:
    logger.warning("Email service not available - emails will not be sent")

# Settings
settings = get_settings()


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize resources on startup, cleanup on shutdown."""
    # Startup
    logger.info("🚀 Starting DOMULEX Backend...")
    
    # Initialize Firebase Auth (optional)
    initialize_firebase()
    
    # Try to initialize Qdrant client (optional - for RAG)
    # Quick initialization - Qdrant will be reconnected lazily if needed
    try:
        # Configure Qdrant client - prefer URL if set
        if settings.qdrant_url:
            # Full URL provided (Qdrant Cloud / Cloud Run)
            logger.info(f"Connecting to Qdrant at {settings.qdrant_url}")
            app.state.qdrant_client = QdrantClient(
                url=settings.qdrant_url,
                api_key=settings.qdrant_api_key if settings.qdrant_api_key else None,
                timeout=5,  # Short timeout for startup
            )
        elif settings.qdrant_use_https:
            # Qdrant Cloud with API Key (legacy)
            app.state.qdrant_client = QdrantClient(
                url=f"https://{settings.qdrant_host}",
                api_key=settings.qdrant_api_key if settings.qdrant_api_key else None,
                timeout=5,
            )
        else:
            app.state.qdrant_client = QdrantClient(
                host=settings.qdrant_host,
                port=settings.qdrant_port,
                timeout=5,
            )
        # Don't test connection at startup - let it fail gracefully later
        logger.info("✅ Qdrant client configured (connection will be tested on first request)")
        
        # Initialize RAG engine with Qdrant
        app.state.rag_engine = RAGEngine(
            qdrant_client=app.state.qdrant_client,
            gemini_api_key=settings.gemini_api_key,
            collection_name=settings.qdrant_collection,
        )
        logger.info("✅ RAG engine initialized with Qdrant")
        
    except Exception as e:
        logger.warning(f"⚠️ Qdrant configuration failed: {e}")
        logger.info("📝 Running in Gemini-only mode (no RAG)")
        app.state.qdrant_client = None
        # Initialize RAG engine without Qdrant (Gemini-only mode)
        app.state.rag_engine = RAGEngine(
            qdrant_client=None,
            gemini_api_key=settings.gemini_api_key,
            collection_name=settings.qdrant_collection,
        )
        logger.info("✅ RAG engine initialized (Gemini-only mode)")
    
    logger.info("✅ DOMULEX Backend ready!")
    
    # Initialize document generator
    if DOCUMENT_GEN_ENABLED:
        app.state.doc_generator = DocumentGenerator(gemini_api_key=settings.gemini_api_key)
        logger.info("✅ Document generator initialized")
    
    # Initialize CRM service
    if CRM_ENABLED:
        app.state.crm_service = CRMService(gemini_api_key=settings.gemini_api_key)
        logger.info("✅ CRM service initialized")
    
    yield
    
    # Shutdown
    logger.info("👋 Shutting down DOMULEX Backend...")


# FastAPI app
app = FastAPI(
    title="DOMULEX Backend",
    description="Legal-Tech RAG System for Real Estate across DE, ES, US",
    version="0.1.0",
    lifespan=lifespan,
)


# CORS middleware for frontend - include all production domains
CORS_ORIGINS = [
    "http://localhost:3000",
    "https://domulex-ai.web.app",
    "https://domulex.ai",
    "https://www.domulex.ai",
]
# Add any additional origins from settings
for origin in settings.get_cors_list():
    if origin not in CORS_ORIGINS:
        CORS_ORIGINS.append(origin)

app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Request logging middleware
app.middleware("http")(RequestLogger(logger))


# Dependency to get RAG engine
def get_rag_engine() -> RAGEngine:
    """Dependency injection for RAG engine."""
    return app.state.rag_engine


# Dependency to get document generator
def get_doc_generator() -> DocumentGenerator:
    """Dependency injection for document generator."""
    if not DOCUMENT_GEN_ENABLED or not hasattr(app.state, 'doc_generator'):
        raise HTTPException(status_code=503, detail="Document generator not available")
    return app.state.doc_generator


# Dependency to get CRM service
def get_crm_service() -> CRMService:
    """Dependency injection for CRM service."""
    if not CRM_ENABLED or not hasattr(app.state, 'crm_service'):
        raise HTTPException(status_code=503, detail="CRM service not available")
    return app.state.crm_service


# Import user service for query counting
from services.user_service import get_user_service

async def count_ai_query(user: FirebaseUser = Depends(get_current_user)):
    """
    Dependency that counts an AI query for the authenticated user.
    Use this dependency on ALL endpoints that use AI/Gemini.
    """
    if user and hasattr(user, 'uid') and user.uid:
        user_service = get_user_service()
        user_service.increment_query_count(user.uid)
        logger.debug(f"Counted AI query for user {user.uid}")
    return user


# === ROUTES ===

@app.get("/")
async def root():
    """Health check."""
    return {
        "status": "ok",
        "service": "DOMULEX Backend",
        "version": "0.1.0",
        "supported_jurisdictions": ["DE", "ES", "US"],
    }


@app.get("/health")
async def health_check():
    """Detailed health check - returns quickly without blocking."""
    # Check Qdrant status without blocking
    if app.state.qdrant_client:
        qdrant_status = "configured (lazy connection)"
    else:
        qdrant_status = "not configured"
    
    return {
        "status": "ok",
        "qdrant": qdrant_status,
        "gemini": "configured" if settings.gemini_api_key else "missing",
    }


@app.get("/health/detailed")
async def health_check_detailed():
    """Detailed health check that actually tests connections (may be slow)."""
    try:
        # Check Qdrant connection
        if app.state.qdrant_client:
            collections = app.state.qdrant_client.get_collections()
            qdrant_status = "ok"
        else:
            qdrant_status = "not configured"
    except Exception as e:
        qdrant_status = f"error: {str(e)}"
    
    return {
        "status": "ok" if qdrant_status == "ok" else "degraded",
        "qdrant": qdrant_status,
        "gemini": "configured" if settings.gemini_api_key else "missing",
    }


# === SESSION MANAGEMENT (Single Device Enforcement) ===

class SessionRegisterRequest(BaseModel):
    """Request to register a new session."""
    session_id: str
    device_info: Optional[str] = None

@app.post("/auth/register-session")
async def register_user_session(
    request: SessionRegisterRequest,
    user: FirebaseUser = Depends(get_current_user)
):
    """
    Register a new session for the authenticated user.
    This invalidates any previous sessions on other devices.
    
    Each user can only be logged in on ONE device at a time.
    """
    result = await register_session(
        uid=user.uid,
        session_id=request.session_id,
        device_info=request.device_info
    )
    
    if result.get("success"):
        return {
            "success": True,
            "message": "Session registered successfully",
            "session_id": request.session_id
        }
    else:
        raise HTTPException(
            status_code=500,
            detail=result.get("error", "Failed to register session")
        )


# === DOCUMENT UPLOAD ===

@app.post("/upload/document")
async def upload_document(
    file: UploadFile = File(...),
    user_id: str = Form(...),
    session_id: Optional[str] = Form(None)
):
    """
    Upload and parse documents (PDF, DOCX, TXT, JPG, PNG).
    
    **Supported Formats:**
    - PDF (text extraction with PyMuPDF)
    - DOCX (text + tables extraction)
    - TXT (plain text)
    - JPG, PNG (OCR mit Tesseract)
    
    **Tier Limits:**
    - Basis: Max 3 Dokumente
    - Professional: Max 20 Dokumente
    - Lawyer Pro: Unbegrenzt
    
    **Returns:**
    - Extracted text preview
    - Cloud Storage URL for later use
    - Metadata (page count, word count, etc.)
    """
    from models.upload import DocumentUploadResponse, UploadedDocument, UploadedDocumentType
    from services.document_parser import document_parser
    import uuid
    
    try:
        logger.info(f"📎 Document upload from user {user_id}: {file.filename}")
        
        # TODO: Check tier limits (implement later with Firestore)
        
        # Read file bytes
        file_bytes = await file.read()
        
        # Parse document
        parse_result = await document_parser.parse_document(file_bytes, file.filename)
        
        if parse_result.error:
            logger.error(f"Document parsing failed: {parse_result.error}")
            return DocumentUploadResponse(
                success=False,
                error=parse_result.error
            )
        
        # TODO: Upload to Cloud Storage (implement later)
        # For now, we'll just return the parsed text
        storage_url = f"temp://documents/{user_id}/{uuid.uuid4()}/{file.filename}"
        
        # Create document metadata
        uploaded_doc = UploadedDocument(
            id=str(uuid.uuid4()),
            user_id=user_id,
            filename=file.filename,
            doc_type=UploadedDocumentType(parse_result.doc_type.value),
            storage_url=storage_url,
            char_count=parse_result.char_count,
            word_count=parse_result.word_count,
            ocr_applied=parse_result.ocr_applied,
            session_id=session_id
        )
        
        # Text preview (first 500 chars) for display
        preview = parse_result.text[:500] + "..." if len(parse_result.text) > 500 else parse_result.text
        
        # Full text for analysis (max 50k chars to prevent memory issues)
        full_text = parse_result.text[:50000] if len(parse_result.text) > 50000 else parse_result.text
        
        logger.info(f"✅ Document uploaded: {uploaded_doc.id} ({parse_result.char_count} chars)")
        
        return DocumentUploadResponse(
            success=True,
            document=uploaded_doc,
            extracted_text_preview=preview,
            extracted_text_full=full_text,  # 📎 Vollständiger Text für Analyse
        )
    
    except Exception as e:
        logger.error(f"Document upload error: {e}", exc_info=True)
        return DocumentUploadResponse(
            success=False,
            error=str(e)
        )


# === SCHRIFTSATZGENERIERUNG (Lawyer Pro only) ===

@app.get("/templates")
async def list_templates():
    """
    Liste aller verfügbaren Dokumentvorlagen.
    
    **Verfügbare Templates:**
    - Klage (Mietrecht)
    - Zahlungsmahnung
    - Kündigung durch Mieter
    - Mängelanzeige
    """
    from services.template_engine import TEMPLATES
    
    return {
        "templates": [
            {
                "id": t.id,
                "name": t.name,
                "category": t.category,
                "description": t.description,
                "icon": t.icon,
                "fields": [
                    {
                        "name": f.name,
                        "label": f.label,
                        "type": f.type,
                        "required": f.required,
                        "placeholder": f.placeholder
                    }
                    for f in t.fields
                ]
            }
            for t in TEMPLATES.values()
        ]
    }


@app.post("/documents/generate-field")
async def generate_field(
    request: dict,
    rag_engine: RAGEngine = Depends(get_rag_engine),
    user: dict = Depends(count_ai_query)
):
    """
    KI-generierung für einzelnes Feld.
    
    **Nur Lawyer Pro Tier!**
    
    Request:
    {
        "template_id": "klage_mietrecht",
        "field_name": "sachverhalt",
        "context_documents": ["Dokument 1 Text...", "Dokument 2 Text..."],
        "user_input": "Optional: Nutzer-Vorgaben"
    }
    """
    from services.template_engine import TEMPLATES, SchriftsatzGenerator
    from services.template_engine import GenerateFieldResponse
    
    try:
        template_id = request.get("template_id")
        field_name = request.get("field_name")
        context_docs = request.get("context_documents", [])
        user_input = request.get("user_input")
        
        if template_id not in TEMPLATES:
            return GenerateFieldResponse(
                success=False,
                field_name=field_name,
                error=f"Template '{template_id}' nicht gefunden"
            )
        
        template = TEMPLATES[template_id]
        field = next((f for f in template.fields if f.name == field_name), None)
        
        if not field:
            return GenerateFieldResponse(
                success=False,
                field_name=field_name,
                error=f"Feld '{field_name}' nicht gefunden"
            )
        
        # Initialisiere Generator
        generator = SchriftsatzGenerator(
            gemini_model=app.state.gemini_model,
            rag_engine=rag_engine
        )
        
        # Generiere Feldinhalt
        content = await generator.generate_field_content(
            field=field,
            context_documents=context_docs,
            user_input=user_input
        )
        
        return GenerateFieldResponse(
            success=True,
            field_name=field_name,
            generated_content=content
        )
    
    except Exception as e:
        logger.error(f"Field generation error: {e}", exc_info=True)
        return GenerateFieldResponse(
            success=False,
            field_name=request.get("field_name", "unknown"),
            error=str(e)
        )


@app.post("/documents/generate")
async def generate_document(
    request: dict,
    rag_engine: RAGEngine = Depends(get_rag_engine),
    user: dict = Depends(count_ai_query)
):
    """
    Generiert vollständiges juristisches Dokument.
    
    **Nur Lawyer Pro Tier!**
    
    Request:
    {
        "template_id": "klage_mietrecht",
        "field_values": {
            "klaeger_name": "Max Mustermann",
            "sachverhalt": "...",
            ...
        },
        "context_documents": ["Dokument 1 Text...", "Dokument 2 Text..."]
    }
    """
    from services.template_engine import TEMPLATES, SchriftsatzGenerator
    from services.template_engine import GenerateDocumentResponse
    
    try:
        template_id = request.get("template_id")
        field_values = request.get("field_values", {})
        context_docs = request.get("context_documents", [])
        
        if template_id not in TEMPLATES:
            return GenerateDocumentResponse(
                success=False,
                error=f"Template '{template_id}' nicht gefunden"
            )
        
        template = TEMPLATES[template_id]
        
        # Initialisiere Generator
        generator = SchriftsatzGenerator(
            gemini_model=app.state.gemini_model,
            rag_engine=rag_engine
        )
        
        # Generiere Dokument
        document = await generator.generate_document(
            template_id=template_id,
            field_values=field_values,
            context_documents=context_docs
        )
        
        return GenerateDocumentResponse(
            success=True,
            template_name=template.name,
            generated_document=document
        )
    
    except Exception as e:
        logger.error(f"Document generation error: {e}", exc_info=True)
        return GenerateDocumentResponse(
            success=False,
            error=str(e)
        )


@app.post("/documents/export")
async def export_document(request: dict):
    """
    Exportiert generiertes Dokument als DOCX oder PDF.
    
    **Nur Lawyer Pro Tier!**
    
    Request:
    {
        "content": "Generierter Dokumenttext...",
        "template_id": "klage_mietrecht",
        "format": "docx" oder "pdf",
        "kanzlei_name": "Optional: Kanzleiname",
        "kanzlei_adresse": "Optional: Adresse"
    }
    
    Returns:
        File download (DOCX oder PDF)
    """
    from services.document_export import document_exporter
    from fastapi.responses import StreamingResponse
    
    try:
        content = request.get("content")
        template_id = request.get("template_id")
        export_format = request.get("format", "docx")
        kanzlei_name = request.get("kanzlei_name")
        kanzlei_adresse = request.get("kanzlei_adresse")
        
        if not content:
            raise HTTPException(status_code=400, detail="Content required")
        
        if export_format == "docx":
            file_bytes = document_exporter.export_docx(
                content=content,
                template_id=template_id,
                kanzlei_name=kanzlei_name,
                kanzlei_adresse=kanzlei_adresse
            )
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{template_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        elif export_format == "pdf":
            file_bytes = document_exporter.export_pdf(
                content=content,
                template_id=template_id,
                kanzlei_name=kanzlei_name,
                kanzlei_adresse=kanzlei_adresse
            )
            media_type = "application/pdf"
            filename = f"{template_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {export_format}")
        
        return StreamingResponse(
            io.BytesIO(file_bytes),
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    
    except Exception as e:
        logger.error(f"Document export error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# =========================================================================
# DOCUMENT ANALYSIS ENDPOINT (KI-gestützte Dokumentenanalyse)
# =========================================================================

class DocumentAnalyzeRequest(BaseModel):
    content: str
    fileName: str

class DocumentAnalyzeResponse(BaseModel):
    suggestedClient: Optional[str] = None
    suggestedCase: Optional[str] = None
    documentType: str
    summary: str
    keywords: List[str] = []

@app.post("/documents/analyze", response_model=DocumentAnalyzeResponse)
async def analyze_document(request: DocumentAnalyzeRequest, user: dict = Depends(count_ai_query)):
    """
    KI-Analyse eines hochgeladenen Dokuments.
    Erkennt automatisch:
    - Dokumenttyp (Mietvertrag, Kündigung, etc.)
    - Mandantenname
    - Aktenzeichen/Sachverhalt
    - Zusammenfassung
    - Schlüsselwörter
    """
    settings = get_settings()
    
    try:
        genai.configure(api_key=settings.GEMINI_API_KEY)
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        prompt = f"""Analysiere dieses juristische Dokument und extrahiere folgende Informationen:

DOKUMENT (Dateiname: {request.fileName}):
{request.content[:8000]}

Antworte NUR mit einem JSON-Objekt in diesem Format:
{{
    "documentType": "Art des Dokuments (z.B. Mietvertrag, Kündigung, Mahnung, Kaufvertrag, Schriftsatz, Urteil)",
    "suggestedClient": "Name der Hauptperson/Mandant falls erkennbar, sonst null",
    "suggestedCase": "Aktenzeichen falls erkennbar, sonst null",
    "summary": "Kurze Zusammenfassung in 1-2 Sätzen auf Deutsch",
    "keywords": ["Schlüsselwort1", "Schlüsselwort2", "Schlüsselwort3"]
}}

Gib NUR das JSON zurück, keine weiteren Erklärungen."""

        response = model.generate_content(prompt)
        response_text = response.text.strip()
        
        # JSON aus Response extrahieren
        import json
        import re
        
        # Entferne mögliche Markdown-Codeblöcke
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0]
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0]
        
        try:
            analysis = json.loads(response_text.strip())
        except json.JSONDecodeError:
            # Fallback: Versuche JSON zu finden
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                analysis = json.loads(json_match.group())
            else:
                raise ValueError("Konnte JSON nicht parsen")
        
        return DocumentAnalyzeResponse(
            documentType=analysis.get("documentType", "Dokument"),
            suggestedClient=analysis.get("suggestedClient"),
            suggestedCase=analysis.get("suggestedCase"),
            summary=analysis.get("summary", "Automatische Analyse abgeschlossen"),
            keywords=analysis.get("keywords", [])
        )
        
    except Exception as e:
        logger.error(f"Document analysis error: {e}", exc_info=True)
        # Fallback ohne KI
        return DocumentAnalyzeResponse(
            documentType="Dokument",
            summary="Automatische Analyse nicht verfügbar",
            keywords=[]
        )


@app.post("/query", response_model=QueryResponse)
async def query_legal_documents(
    request: QueryRequest,
    rag_engine: RAGEngine = Depends(get_rag_engine),
):
    """
    Query legal documents with RAG.
    
    **Tier Limits:**
    - Test: 3 queries total (account deleted after 6 months if not upgraded)
    - Basis: 25 queries/month
    - Professional: 250 queries/month
    - Lawyer Pro: Unlimited
    
    **Critical:** The target_jurisdiction filter prevents cross-jurisdictional hallucinations.
    - Only queries German law (DE jurisdiction)
    
    **Example Request:**
    ```json
    {
        "query": "Was sind meine Rechte bei Schimmel in der Wohnung?",
        "target_jurisdiction": "DE",
        "user_role": "TENANT",
        "user_language": "de",
        "user_id": "firebase_uid_here",
        "user_tier": "free"
    }
    ```
    
    **Response:**
    - Answer in user's language (de)
    - Sources from German law only
    """
    try:
        # Check if user has provided ID for quota tracking
        if request.user_id:
            from services.user_service import get_user_service
            user_service = get_user_service()
            
            # Check query limit before processing
            can_query, queries_used, queries_limit = user_service.check_query_limit(request.user_id)
            
            if not can_query:
                logger.warning(f"User {request.user_id} exceeded query limit: {queries_used}/{queries_limit}")
                raise HTTPException(
                    status_code=429,
                    detail=f"Monatliches Anfrage-Limit erreicht ({queries_used}/{queries_limit}). Bitte upgraden Sie Ihren Plan."
                )
            
            logger.info(f"Query from user {request.user_id}: {queries_used + 1}/{queries_limit}")
        
        # Prepare uploaded documents for RAG engine
        uploaded_docs = None
        if request.uploaded_documents:
            uploaded_docs = [
                {"document_id": doc.document_id, "text": doc.text}
                for doc in request.uploaded_documents
            ]
            logger.info(f"📎 Query includes {len(uploaded_docs)} uploaded documents")
        
        response = await rag_engine.query(
            user_query=request.query,
            target_jurisdiction=request.target_jurisdiction,
            user_role=request.user_role,
            user_language=request.user_language,
            sub_jurisdiction=request.sub_jurisdiction,
            source_filter=request.source_filter,
            gerichtsebene_filter=request.gerichtsebene_filter,
            use_public_sources=request.use_public_sources,  # 🔑 Öffentliche Quellen
            uploaded_documents=uploaded_docs,  # 📎 Hochgeladene Dokumente
        )
        
        # Query count is now handled by count_ai_query dependency
        # (user_id is passed via request for anonymous queries only)
        if request.user_id and not hasattr(request, '_counted'):
            get_user_service().increment_query_count(request.user_id)
        
        return response
    
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error processing query: {str(e)}",
        )


# === SUPPORT CHAT ENDPOINT ===

class SupportRequest(BaseModel):
    """Request for support chat."""
    query: str
    user_language: str = "de"

class SupportResponse(BaseModel):
    """Response from support chat."""
    answer: str
    category: str = "general"

@app.post("/support", response_model=SupportResponse)
async def support_chat(request: SupportRequest):
    """
    Customer support chat - answers questions about Domulex.ai platform.
    
    This does NOT use RAG or legal documents. It uses Gemini directly with
    a support persona that knows about the platform, pricing, and features.
    """
    try:
        from rag.personas import get_support_prompt
        
        # Configure Gemini
        genai.configure(api_key=settings.gemini_api_key)
        model = genai.GenerativeModel('gemini-2.0-flash')
        
        # Build prompt with support persona
        system_prompt = get_support_prompt(request.user_language)
        full_prompt = f"{system_prompt}\n\nKundenanfrage: {request.query}"
        
        # Generate response
        response = model.generate_content(full_prompt)
        answer = response.text
        
        # Detect category
        category = "general"
        query_lower = request.query.lower()
        if any(w in query_lower for w in ["kündigen", "abo", "abonnement", "tarif"]):
            category = "subscription"
        elif any(w in query_lower for w in ["zahlung", "bezahlen", "rechnung", "preis"]):
            category = "billing"
        elif any(w in query_lower for w in ["passwort", "login", "konto", "registrier"]):
            category = "account"
        elif any(w in query_lower for w in ["funktion", "wie", "nutzen", "anfragen"]):
            category = "usage"
        
        return SupportResponse(answer=answer, category=category)
    
    except Exception as e:
        logger.error(f"Support chat error: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Error processing support request: {str(e)}",
        )


class SupportContactRequest(BaseModel):
    """Request for support contact form."""
    email: str
    name: str = "Nicht angegeben"
    message: str
    chat_history: str = ""


@app.post("/support/contact")
async def support_contact(request: SupportContactRequest):
    """
    Handle support contact form submissions.
    Sends email to support team with customer message and chat history.
    Uses unified email templates for consistent branding.
    """
    try:
        if not EMAIL_CLIENT_ENABLED:
            raise HTTPException(status_code=503, detail="Email service not available")
        
        email_client.reload_config()
        
        # Import unified templates
        from services.email_templates import (
            get_support_contact_confirmation_email,
            get_support_contact_internal_email
        )
        
        # Get internal email for support team
        internal_email = get_support_contact_internal_email(
            user_name=request.name,
            user_email=request.email,
            user_message=request.message,
            chat_history=request.chat_history
        )
        
        # Send to support email
        result = email_client.send_email(
            to="kontakt@domulex.ai",
            subject=internal_email["subject"],
            body_html=internal_email["html"],
            body_text=internal_email["text"],
            reply_to=request.email
        )
        
        if result.get("success"):
            logger.info(f"Support contact email sent from {request.email}")
            
            # Send confirmation email to user using unified template
            try:
                confirmation_email = get_support_contact_confirmation_email(
                    user_name=request.name,
                    user_message=request.message
                )
                
                email_client.send_email(
                    to=request.email,
                    subject=confirmation_email["subject"],
                    body_html=confirmation_email["html"],
                    body_text=confirmation_email["text"]
                )
                logger.info(f"Confirmation email sent to {request.email}")
            except Exception as confirm_err:
                logger.warning(f"Could not send confirmation email: {confirm_err}")
            
            return {"success": True, "message": "Ihre Nachricht wurde erfolgreich gesendet."}
        else:
            raise HTTPException(status_code=500, detail=result.get("error", "Failed to send email"))
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Support contact error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/ingest/run", response_model=IngestionResponse)
async def run_ingestion(
    request: IngestionRequest,
    rag_engine: RAGEngine = Depends(get_rag_engine),
):
    """
    Trigger data ingestion for a specific jurisdiction.
    
    **Security:** This endpoint should be protected (add auth in production).
    
    **Example Request:**
    ```json
    {
        "jurisdiction": "DE",
        "force_refresh": false,
        "max_documents": 100
    }
    ```
    
    **Workflow:**
    1. Get scraper for jurisdiction (GermanScraper, SpanishScraper, USScraper)
    2. Fetch documents from official sources
    3. Index into Qdrant with embeddings
    """
    start_time = time.time()
    errors = []
    
    try:
        # Get appropriate scraper
        scraper = ScraperFactory.get_scraper(request.jurisdiction)
        
        # Validate source accessibility
        if not await scraper.validate_source():
            raise HTTPException(
                status_code=503,
                detail=f"Data source for {request.jurisdiction} is not accessible",
            )
        
        # Fetch documents
        documents = await scraper.fetch(max_documents=request.max_documents)
        
        if not documents:
            return IngestionResponse(
                status="completed",
                documents_processed=0,
                errors=["No documents found"],
                duration_seconds=time.time() - start_time,
            )
        
        # Index into Qdrant
        indexed_count = await rag_engine.index_documents(documents)
        
        duration = time.time() - start_time
        
        return IngestionResponse(
            status="completed",
            documents_processed=indexed_count,
            errors=errors,
            duration_seconds=duration,
        )
    
    except Exception as e:
        duration = time.time() - start_time
        raise HTTPException(
            status_code=500,
            detail=f"Ingestion failed: {str(e)}",
        )


@app.get("/jurisdictions")
async def list_jurisdictions():
    """List all supported jurisdictions with metadata."""
    return {
        "jurisdictions": [
            {
                "code": "DE",
                "name": "Germany",
                "legal_system": "Civil Law",
                "languages": ["de"],
                "sub_jurisdictions": ["Bayern", "NRW", "Berlin"],
                "sources": [
                    "gesetze-im-internet.de",
                    "rechtsprechung-im-internet.de",
                ],
            },
            {
                "code": "ES",
                "name": "Spain",
                "legal_system": "Civil Law",
                "languages": ["es"],
                "sub_jurisdictions": ["Madrid", "Catalunya", "Baleares"],
                "sources": ["BOE.es", "Código Civil"],
            },
            {
                "code": "US",
                "name": "United States",
                "legal_system": "Common Law",
                "languages": ["en"],
                "sub_jurisdictions": ["Florida", "New York", "California"],
                "sources": ["CourtListener", "congress.gov"],
            },
        ]
    }


@app.get("/stats")
async def get_statistics():
    """Get system statistics."""
    try:
        collection_info = app.state.qdrant_client.get_collection(
            settings.qdrant_collection
        )
        
        return {
            "total_documents": collection_info.points_count,
            "collection_name": settings.qdrant_collection,
            "vector_size": collection_info.config.params.vectors.size,
        }
    
    except Exception as e:
        return {
            "error": str(e),
            "total_documents": 0,
        }


# Helper function for text extraction using Gemini Vision
async def extract_text_with_gemini(file_bytes: bytes, filename: str) -> str:
    """Extract text from images or documents using Gemini Vision API."""
    try:
        import base64
        genai.configure(api_key=settings.gemini_api_key)
        model = genai.GenerativeModel('gemini-2.0-flash')
        
        # Determine MIME type
        ext = filename.lower().split('.')[-1]
        mime_types = {
            'jpg': 'image/jpeg', 'jpeg': 'image/jpeg', 'png': 'image/png',
            'gif': 'image/gif', 'webp': 'image/webp', 'bmp': 'image/bmp',
            'tiff': 'image/tiff', 'pdf': 'application/pdf',
            'doc': 'application/msword', 'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        }
        mime_type = mime_types.get(ext, 'application/octet-stream')
        
        # For images, use vision capabilities
        if ext in ['jpg', 'jpeg', 'png', 'gif', 'webp', 'bmp', 'tiff']:
            image_data = base64.b64encode(file_bytes).decode('utf-8')
            response = model.generate_content([
                {"mime_type": mime_type, "data": image_data},
                "Extrahiere den gesamten Text aus diesem Dokument/Bild. Gib nur den extrahierten Text zurück, keine Erklärungen."
            ])
            return response.text
        else:
            # For other documents, try to describe content
            file_data = base64.b64encode(file_bytes).decode('utf-8')
            response = model.generate_content([
                {"mime_type": mime_type, "data": file_data},
                "Extrahiere den gesamten Text aus diesem Dokument. Gib nur den extrahierten Text zurück."
            ])
            return response.text
    except Exception as e:
        logger.error(f"Gemini text extraction failed: {e}")
        return ""


@app.post("/analyze_contract", response_model=ContractAnalysis)
async def analyze_contract(
    file: UploadFile = File(..., description="PDF contract file"),
    jurisdiction: str = Form(..., description="Target jurisdiction (DE, ES, US)"),
    user_role: str = Form(..., description="User role (TENANT, LANDLORD, etc.)"),
    user_tier: str = Form(..., description="User subscription tier (free, mieter_plus, professional, lawyer)"),
    rag_engine: RAGEngine = Depends(get_rag_engine),
    user: dict = Depends(count_ai_query),
):
    """
    Analyze a PDF contract and compare clauses against legal standards.
    
    **REQUIRES: Professional or Lawyer tier**
    
    **Workflow:**
    1. Extract text from PDF
    2. Identify key clauses using Gemini
    3. For each clause, search Qdrant for relevant laws
    4. Compare contract clause vs. legal standard
    5. Generate risk assessment (RED/YELLOW/GREEN)
    
    **Example Request:**
    ```bash
    curl -X POST http://localhost:8000/analyze_contract \
      -F "file=@lease_agreement.pdf" \
      -F "jurisdiction=DE" \
      -F "user_role=TENANT" \
      -F "user_tier=professional"
    ```
    """
    try:
        # Check tier access - PDF analysis only for Professional and Lawyer
        if user_tier.lower() not in ['professional', 'lawyer']:
            raise HTTPException(
                status_code=403,
                detail="PDF-Vertragsanalyse ist nur im Professional- oder Lawyer-Tarif verfügbar. Bitte upgraden Sie Ihren Tarif."
            )
        # Validate inputs
        try:
            jurisdiction_enum = Jurisdiction(jurisdiction)
            user_role_enum = UserRole(user_role)
        except ValueError as e:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid jurisdiction or role: {str(e)}"
            )
        
        # Validate file type - Lawyers get all formats, others only PDF
        filename_lower = file.filename.lower()
        is_lawyer = user_tier.lower() == 'lawyer'
        
        # Allowed extensions by tier
        pdf_extensions = ['.pdf']
        lawyer_extensions = ['.pdf', '.doc', '.docx', '.xls', '.xlsx', '.jpg', '.jpeg', '.png', 
                            '.gif', '.webp', '.bmp', '.tiff', '.txt', '.csv', '.html', '.htm',
                            '.rtf', '.eml', '.msg']
        
        allowed_extensions = lawyer_extensions if is_lawyer else pdf_extensions
        file_extension = '.' + filename_lower.split('.')[-1] if '.' in filename_lower else ''
        
        if file_extension not in allowed_extensions:
            if is_lawyer:
                raise HTTPException(
                    status_code=400,
                    detail="Nicht unterstütztes Dateiformat. Erlaubt: PDF, Word, Excel, Bilder, Text, HTML, RTF, E-Mail"
                )
            else:
                raise HTTPException(
                    status_code=400,
                    detail="Nur PDF-Dateien werden unterstützt"
                )
        
        # Read file bytes
        file_bytes = await file.read()
        
        # File size limits: 10MB for regular, 25MB for lawyers
        max_size = 25 * 1024 * 1024 if is_lawyer else 10 * 1024 * 1024
        if len(file_bytes) > max_size:
            raise HTTPException(
                status_code=400,
                detail=f"Datei zu groß (max. {25 if is_lawyer else 10}MB)"
            )
        
        # Step A: Extract text based on file type
        contract_text = ""
        
        if file_extension == '.pdf':
            contract_text = PDFParser.extract_text_from_pdf(file_bytes)
        elif file_extension in ['.doc', '.docx']:
            # Word documents
            try:
                import docx
                import io
                doc_file = docx.Document(io.BytesIO(file_bytes))
                contract_text = '\n'.join([para.text for para in doc_file.paragraphs])
            except ImportError:
                # Fallback: Use Gemini Vision for document
                contract_text = await extract_text_with_gemini(file_bytes, filename_lower)
        elif file_extension in ['.xls', '.xlsx']:
            # Excel files
            try:
                import pandas as pd
                import io
                df = pd.read_excel(io.BytesIO(file_bytes))
                contract_text = df.to_string()
            except ImportError:
                contract_text = await extract_text_with_gemini(file_bytes, filename_lower)
        elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp', '.tiff']:
            # Images - use OCR via Gemini Vision
            contract_text = await extract_text_with_gemini(file_bytes, filename_lower)
        elif file_extension in ['.txt', '.csv', '.html', '.htm', '.rtf']:
            # Text files
            try:
                contract_text = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                contract_text = file_bytes.decode('latin-1', errors='ignore')
        elif file_extension in ['.eml', '.msg']:
            # Email files - extract body
            try:
                import email
                if file_extension == '.eml':
                    msg = email.message_from_bytes(file_bytes)
                    if msg.is_multipart():
                        for part in msg.walk():
                            if part.get_content_type() == 'text/plain':
                                contract_text += part.get_payload(decode=True).decode('utf-8', errors='ignore')
                    else:
                        contract_text = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
                else:
                    # .msg files need extract-msg library
                    contract_text = await extract_text_with_gemini(file_bytes, filename_lower)
            except Exception:
                contract_text = await extract_text_with_gemini(file_bytes, filename_lower)
        else:
            # Fallback to Gemini
            contract_text = await extract_text_with_gemini(file_bytes, filename_lower)
        
        if not contract_text or len(contract_text.strip()) < 50:
            raise HTTPException(
                status_code=422,
                detail="Konnte keinen Text aus der Datei extrahieren. Bitte laden Sie eine lesbare Datei hoch."
            )
        
        # Initialize Gemini model (use gemini-2.0-flash for better availability)
        genai.configure(api_key=settings.gemini_api_key)
        gemini_model = genai.GenerativeModel('gemini-2.0-flash')
        
        # Step B: Identify key clauses
        identified_clauses = PDFParser.identify_key_clauses(contract_text, gemini_model)
        
        if not identified_clauses:
            raise HTTPException(
                status_code=422,
                detail="Could not identify any key clauses in the contract"
            )
        
        # Step C & D: Analyze each clause
        clause_analyses = []
        red_count = 0
        yellow_count = 0
        green_count = 0
        
        for clause_data in identified_clauses:
            clause_type = clause_data.get("type", "Unknown")
            clause_text = clause_data.get("text", "")
            
            # Search Qdrant for relevant law
            search_query = f"{clause_type} law {jurisdiction_enum.value} real estate lease"
            
            # 🔑 Quellenfilter basierend auf User-Tier
            # Basis/Professional: Nur verlässliche Quellen (Gesetze + Höchstgerichte)
            # Lawyer: Alle Quellen
            if user_tier.lower() == 'lawyer':
                source_filter = None  # Alle Quellen
                gerichtsebene_filter = None
            else:
                source_filter = ['GESETZ', 'URTEIL', 'VERWALTUNG']  # Nur Gesetze, Urteile, Verwaltung
                gerichtsebene_filter = ['EuGH', 'BGH', 'BFH']  # Nur Höchstgerichte - verlässlich
            
            search_results = await rag_engine.query(
                user_query=search_query,
                target_jurisdiction=jurisdiction_enum,
                user_role=user_role_enum,
                user_language="en",
                sub_jurisdiction=None,
                source_filter=source_filter,
                gerichtsebene_filter=gerichtsebene_filter,
            )
            
            # Extract legal context from search results
            legal_context = ""
            source_title = None
            source_url = None
            
            if search_results.sources:
                top_source = search_results.sources[0]
                legal_context = top_source.content_original
                source_title = top_source.title
                source_url = str(top_source.source_url)
            else:
                legal_context = "Keine spezifische Rechtsgrundlage in der Datenbank gefunden. Die Analyse basiert auf allgemeinem deutschem Immobilienrecht."
            
            # Compare clause with law
            comparison_result = PDFParser.compare_clause_with_law(
                clause_type=clause_type,
                clause_text=clause_text,
                legal_context=legal_context,
                user_role=user_role_enum,
                jurisdiction=jurisdiction_enum,
                gemini_model=gemini_model,
            )
            
            # Create clause analysis
            risk_level = RiskLevel(comparison_result.get("risk_level", "YELLOW"))
            
            clause_analysis = ClauseAnalysis(
                clause_type=clause_type,
                clause_text=clause_text,
                risk_level=risk_level,
                legal_standard=legal_context[:300] + "..." if len(legal_context) > 300 else legal_context,
                comparison=comparison_result.get("comparison", "No comparison available"),
                recommendation=comparison_result.get("recommendation"),
                source_title=source_title,
                source_url=source_url,
            )
            
            clause_analyses.append(clause_analysis)
            
            # Count flags
            if risk_level == RiskLevel.RED:
                red_count += 1
            elif risk_level == RiskLevel.YELLOW:
                yellow_count += 1
            else:
                green_count += 1
        
        # Determine overall risk
        if red_count > 0:
            overall_risk = RiskLevel.RED
        elif yellow_count > red_count:
            overall_risk = RiskLevel.YELLOW
        else:
            overall_risk = RiskLevel.GREEN
        
        # Generate summary (German)
        summary = f"{len(clause_analyses)} Klauseln analysiert. "
        summary += f"{red_count} kritische Probleme (ROT), "
        summary += f"{yellow_count} potenzielle Bedenken (GELB), "
        summary += f"{green_count} unbedenkliche Klauseln (GRÜN). "
        
        if red_count > 0:
            summary += "⚠️ EMPFEHLUNG: Lassen Sie den Vertrag vor Unterzeichnung von einem Rechtsanwalt prüfen."
        elif yellow_count > 0:
            summary += "⚠️ EMPFEHLUNG: Prüfen Sie die markierten Klauseln sorgfältig."
        else:
            summary += "✅ Der Vertrag erscheint grundsätzlich unbedenklich. Eine anwaltliche Prüfung ist dennoch empfehlenswert."
        
        # Create analysis result
        analysis = ContractAnalysis(
            contract_name=file.filename,
            jurisdiction=jurisdiction_enum,
            user_role=user_role_enum,
            clauses=clause_analyses,
            overall_risk=overall_risk,
            summary=summary,
            total_clauses_analyzed=len(clause_analyses),
            red_flags=red_count,
            yellow_flags=yellow_count,
            green_flags=green_count,
        )
        
        return analysis
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Contract analysis failed: {str(e)}"
        )


# ============ LAWYER-ONLY ENDPOINTS ============

class AlternativeClauseRequest(BaseModel):
    clause_text: str
    clause_type: str
    user_role: str = "TENANT"
    risk_level: str = "YELLOW"

class AlternativeClauseResponse(BaseModel):
    alternatives: List[str]
    explanation: str

@app.post("/generate_alternative_clause", response_model=AlternativeClauseResponse)
async def generate_alternative_clause(request: AlternativeClauseRequest, user: dict = Depends(count_ai_query)):
    """
    Generate alternative clause suggestions for lawyers.
    
    **REQUIRES: Lawyer tier**
    
    Takes a problematic clause and generates 2-3 alternative formulations
    that are more favorable for the client's position.
    """
    try:
        genai.configure(api_key=settings.gemini_api_key)
        gemini_model = genai.GenerativeModel('gemini-2.0-flash')
        
        role_name = {
            'TENANT': 'Mieter',
            'LANDLORD': 'Vermieter',
            'BUYER': 'Käufer',
            'SELLER': 'Verkäufer'
        }.get(request.user_role, 'Mandant')
        
        prompt = f"""Du bist ein erfahrener deutscher Rechtsanwalt für Immobilienrecht.

AUFGABE: Generiere 2-3 alternative Klauselformulierungen für die folgende problematische Vertragsklausel.

ORIGINALKLAUSEL ({request.clause_type}):
"{request.clause_text}"

RISIKOSTUFE: {request.risk_level}
MANDANTENROLLE: {role_name}

ANFORDERUNGEN:
1. Jede Alternative muss rechtlich zulässig nach deutschem Recht sein
2. Die Alternativen sollen die Position des {role_name}s stärken
3. Formuliere präzise, rechtssicher und verständlich
4. Behalte die wesentliche Vertragsintention bei, aber entschärfe problematische Aspekte

Antworte im folgenden JSON-Format:
{{
  "alternatives": [
    "Alternative Klausel 1...",
    "Alternative Klausel 2...",
    "Alternative Klausel 3..."
  ],
  "explanation": "Kurze Erklärung warum diese Alternativen besser sind..."
}}"""

        response = gemini_model.generate_content(prompt)
        response_text = response.text.strip()
        
        # Parse JSON from response
        import json
        import re
        
        # Extract JSON from markdown code blocks if present
        json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response_text, re.DOTALL)
        if json_match:
            response_text = json_match.group(1)
        
        try:
            result = json.loads(response_text)
            return AlternativeClauseResponse(
                alternatives=result.get("alternatives", []),
                explanation=result.get("explanation", "")
            )
        except json.JSONDecodeError:
            # Fallback: extract alternatives manually
            return AlternativeClauseResponse(
                alternatives=[response_text],
                explanation="Automatisch generierte Alternative"
            )
            
    except Exception as e:
        logger.error(f"Alternative clause generation failed: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Klauselgenerierung fehlgeschlagen: {str(e)}"
        )


class ExportReportRequest(BaseModel):
    analysis: dict
    client_reference: Optional[str] = None
    user_role: str = "TENANT"

from fastapi.responses import StreamingResponse
import io

@app.post("/export_contract_report")
async def export_contract_report(request: ExportReportRequest):
    """
    Export contract analysis as a professional PDF legal opinion (Gutachten).
    
    **REQUIRES: Lawyer tier**
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        from datetime import datetime
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                                leftMargin=2.5*cm, rightMargin=2.5*cm,
                                topMargin=2.5*cm, bottomMargin=2.5*cm)
        
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='TitleDE', 
                                  fontSize=18, 
                                  spaceAfter=20, 
                                  alignment=TA_CENTER,
                                  fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='Heading2DE', 
                                  fontSize=14, 
                                  spaceAfter=10, 
                                  spaceBefore=15,
                                  fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='BodyDE', 
                                  fontSize=11, 
                                  spaceAfter=8,
                                  alignment=TA_JUSTIFY,
                                  leading=14))
        
        story = []
        
        # Header
        story.append(Paragraph("RECHTSGUTACHTEN", styles['TitleDE']))
        story.append(Paragraph("Vertragsanalyse", styles['Heading2DE']))
        story.append(Spacer(1, 20))
        
        # Meta info
        analysis = request.analysis
        role_name = {'TENANT': 'Mieter', 'LANDLORD': 'Vermieter', 
                     'BUYER': 'Käufer', 'SELLER': 'Verkäufer'}.get(request.user_role, 'Mandant')
        
        meta_data = [
            ['Vertrag:', analysis.get('contract_name', 'Unbekannt')],
            ['Datum:', datetime.now().strftime('%d.%m.%Y')],
            ['Mandantenrolle:', role_name],
        ]
        if request.client_reference:
            meta_data.append(['Aktenzeichen:', request.client_reference])
        
        meta_table = Table(meta_data, colWidths=[4*cm, 10*cm])
        meta_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 20))
        
        # Overall Risk
        risk = analysis.get('overall_risk', 'YELLOW')
        risk_text = {'RED': 'KRITISCH - Dringender Handlungsbedarf', 
                     'YELLOW': 'PRÜFENSWERT - Einzelne Punkte beachten',
                     'GREEN': 'UNBEDENKLICH - Vertrag entspricht Standards'}.get(risk, 'UNBEKANNT')
        
        story.append(Paragraph(f"Gesamtbewertung: {risk_text}", styles['Heading2DE']))
        story.append(Paragraph(analysis.get('summary', ''), styles['BodyDE']))
        story.append(Spacer(1, 15))
        
        # Statistics
        story.append(Paragraph("Zusammenfassung", styles['Heading2DE']))
        stats_data = [
            ['Kategorie', 'Anzahl'],
            ['Analysierte Klauseln', str(analysis.get('total_clauses_analyzed', 0))],
            ['Kritische Probleme (ROT)', str(analysis.get('red_flags', 0))],
            ['Potenzielle Bedenken (GELB)', str(analysis.get('yellow_flags', 0))],
            ['Unbedenklich (GRÜN)', str(analysis.get('green_flags', 0))],
        ]
        stats_table = Table(stats_data, colWidths=[10*cm, 4*cm])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e3a5f')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(stats_table)
        story.append(Spacer(1, 20))
        
        # Detailed Clause Analysis
        story.append(Paragraph("Detailanalyse der Klauseln", styles['Heading2DE']))
        
        clauses = analysis.get('clauses', [])
        for idx, clause in enumerate(clauses, 1):
            risk_color = {'RED': colors.HexColor('#dc2626'), 
                          'YELLOW': colors.HexColor('#ca8a04'),
                          'GREEN': colors.HexColor('#16a34a')}.get(clause.get('risk_level'), colors.grey)
            
            story.append(Spacer(1, 10))
            story.append(Paragraph(f"{idx}. {clause.get('clause_type', 'Klausel')}", styles['Heading2DE']))
            
            # Clause text box
            clause_text = clause.get('clause_text', '')[:500]
            story.append(Paragraph(f"<i>\"{clause_text}...\"</i>", styles['BodyDE']))
            
            # Risk and comparison
            risk_label = {'RED': 'KRITISCH', 'YELLOW': 'PRÜFENSWERT', 'GREEN': 'OK'}.get(clause.get('risk_level'), '?')
            story.append(Paragraph(f"<b>Bewertung:</b> {risk_label}", styles['BodyDE']))
            story.append(Paragraph(f"<b>Analyse:</b> {clause.get('comparison', '')}", styles['BodyDE']))
            
            if clause.get('recommendation'):
                story.append(Paragraph(f"<b>Empfehlung:</b> {clause.get('recommendation')}", styles['BodyDE']))
        
        # Footer / Disclaimer
        story.append(Spacer(1, 30))
        story.append(Paragraph("Rechtlicher Hinweis", styles['Heading2DE']))
        story.append(Paragraph(
            "Dieses Gutachten wurde mit Unterstützung künstlicher Intelligenz erstellt. "
            "Es ersetzt keine individuelle Rechtsberatung. Die Analyse basiert auf den zum "
            "Zeitpunkt der Erstellung verfügbaren Informationen und der aktuellen Rechtslage. "
            "Für verbindliche rechtliche Einschätzungen konsultieren Sie bitte einen Fachanwalt.",
            styles['BodyDE']
        ))
        story.append(Spacer(1, 15))
        story.append(Paragraph(f"Erstellt am {datetime.now().strftime('%d.%m.%Y um %H:%M Uhr')} via domulex.ai", 
                               ParagraphStyle('Footer', fontSize=9, textColor=colors.grey)))
        
        doc.build(story)
        buffer.seek(0)
        
        filename = f"Rechtsgutachten_{analysis.get('contract_name', 'Vertrag')}_{datetime.now().strftime('%Y%m%d')}.pdf"
        
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="PDF-Export nicht verfügbar. Bitte installieren Sie reportlab."
        )
    except Exception as e:
        logger.error(f"PDF export failed: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"PDF-Export fehlgeschlagen: {str(e)}"
        )


@app.post("/resolve_conflict", response_model=ConflictResponse)
async def resolve_conflict(
    request: ConflictRequest,
    rag_engine: RAGEngine = Depends(get_rag_engine),
):
    """
    Analyze a legal dispute from both perspectives using neutral mediation.
    
    **Workflow:**
    1. Combine both statements and embed them
    2. Search Qdrant for relevant case law and precedents
    3. Use MEDIATOR persona to analyze from both sides
    4. Calculate success probabilities based on legal precedent
    5. Provide neutral recommendation
    
    **Example Request:**
    ```bash
    curl -X POST http://localhost:8000/resolve_conflict \\
      -H "Content-Type: application/json" \\
      -d '{
        "party_a_statement": "I gave 30 days notice but tenant won't leave",
        "party_b_statement": "Landlord didn't fix broken AC for 3 months",
        "jurisdiction": "US",
        "party_a_label": "Landlord",
        "party_b_label": "Tenant"
      }'
    ```
    """
    try:
        # Initialize Gemini
        genai.configure(api_key=settings.gemini_api_key)
        gemini_model = genai.GenerativeModel('gemini-1.5-pro')
        
        # Combine statements for embedding
        combined_query = f"""
        Dispute Context:
        {request.party_a_label}: {request.party_a_statement}
        {request.party_b_label}: {request.party_b_statement}
        """
        
        # Search for relevant legal precedents
        search_results = await rag_engine.query(
            user_query=combined_query,
            target_jurisdiction=request.jurisdiction,
            user_role=UserRole.MEDIATOR,
            user_language=request.user_language,
            sub_jurisdiction=request.sub_jurisdiction,
        )
        
        # Prepare legal context from sources
        legal_context = ""
        all_sources = search_results.sources
        
        if all_sources:
            legal_context = "\n\n---\n\n".join([
                f"**{doc.title}**\n{doc.content_original[:500]}..."
                for doc in all_sources[:5]  # Top 5 most relevant
            ])
        else:
            legal_context = "No specific precedent found in database."
        
        # Get mediator prompt
        mediator_prompt = get_mediator_prompt(
            jurisdiction=request.jurisdiction,
            user_language=request.user_language,
        )
        
        # Construct analysis prompt
        analysis_prompt = f"""
{mediator_prompt}

---

**DISPUTE DETAILS:**

**{request.party_a_label}'s Statement:**
{request.party_a_statement}

**{request.party_b_label}'s Statement:**
{request.party_b_statement}

**Relevant {request.jurisdiction.value} Legal Precedents:**
{legal_context}

---

**Provide your neutral mediation analysis following the protocol above.**

IMPORTANT: Return your response in the following JSON structure:
```json
{{
  "dispute_summary": "Brief 1-2 sentence summary",
  "party_a_arguments": "Legal arguments supporting {request.party_a_label}",
  "party_a_strength": "Weak/Moderate/Strong with reasoning",
  "party_b_arguments": "Legal arguments supporting {request.party_b_label}",
  "party_b_strength": "Weak/Moderate/Strong with reasoning",
  "neutral_assessment": "Objective analysis of legal situation",
  "success_probability_a": 45,
  "success_probability_b": 55,
  "settlement_likelihood": 75,
  "recommendation": "Specific recommendation with reasoning"
}}
```
"""
        
        # Generate mediation analysis
        response = gemini_model.generate_content(analysis_prompt)
        response_text = response.text.strip()
        
        # Parse JSON response
        import json
        if "```json" in response_text:
            response_text = response_text.split("```json")[1].split("```")[0].strip()
        elif "```" in response_text:
            response_text = response_text.split("```")[1].split("```")[0].strip()
        
        analysis_data = json.loads(response_text)
        
        # Create party analyses
        party_a_sources = [doc for doc in all_sources if request.party_a_label.lower() in doc.content_original.lower()][:3]
        party_b_sources = [doc for doc in all_sources if request.party_b_label.lower() in doc.content_original.lower()][:3]
        
        # If no specific sources, use top sources for both
        if not party_a_sources:
            party_a_sources = all_sources[:2]
        if not party_b_sources:
            party_b_sources = all_sources[:2]
        
        party_a_analysis = ConflictAnalysis(
            party_label=request.party_a_label,
            legal_arguments=analysis_data.get("party_a_arguments", "No arguments found"),
            supporting_sources=party_a_sources,
            strength_assessment=analysis_data.get("party_a_strength", "Moderate"),
        )
        
        party_b_analysis = ConflictAnalysis(
            party_label=request.party_b_label,
            legal_arguments=analysis_data.get("party_b_arguments", "No arguments found"),
            supporting_sources=party_b_sources,
            strength_assessment=analysis_data.get("party_b_strength", "Moderate"),
        )
        
        # Create conflict response
        conflict_response = ConflictResponse(
            dispute_summary=analysis_data.get("dispute_summary", "Dispute analysis"),
            party_a_analysis=party_a_analysis,
            party_b_analysis=party_b_analysis,
            neutral_assessment=analysis_data.get("neutral_assessment", ""),
            success_probability_a=float(analysis_data.get("success_probability_a", 50)),
            success_probability_b=float(analysis_data.get("success_probability_b", 50)),
            settlement_likelihood=float(analysis_data.get("settlement_likelihood", 50)),
            recommendation=analysis_data.get("recommendation", ""),
            jurisdiction=request.jurisdiction,
        )
        
        return conflict_response
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Conflict resolution analysis failed: {str(e)}"
        )


# === STRIPE PAYMENT ENDPOINTS ===

@app.post("/stripe/create-checkout-session")
async def create_stripe_checkout(
    request: dict,
    current_user = Depends(get_current_user)
):
    """
    Create a Stripe Checkout Session for subscription
    
    **Request Body:**
    - tier: Subscription tier (free, mieter_plus, professional, lawyer)
    - success_url: URL to redirect after successful payment
    - cancel_url: URL to redirect if payment is cancelled
    
    **Returns:**
    - checkout_url: URL to redirect user to Stripe Checkout
    - session_id: Stripe Checkout Session ID
    """
    if not STRIPE_ENABLED:
        raise HTTPException(status_code=503, detail="Stripe payment service not configured")
    
    try:
        tier = request.get("tier", "").lower()
        success_url = request.get("success_url", "")
        cancel_url = request.get("cancel_url", "")
        
        # Map frontend tier IDs to Stripe tier names
        tier_mapping = {
            "free": "FREE",
            "basis": "TENANT",
            "mieter_plus": "TENANT",  # Legacy support
            "professional": "PRO",
            "lawyer": "LAWYER"
        }
        
        if tier not in tier_mapping:
            raise HTTPException(status_code=400, detail=f"Invalid tier: {tier}")
        
        stripe_tier = tier_mapping[tier]
        
        if stripe_tier == "FREE":
            raise HTTPException(status_code=400, detail="Cannot create checkout for free tier")
        
        # Get user info from Firebase token
        user_id = current_user.uid
        user_email = current_user.email
        
        if not user_email:
            raise HTTPException(status_code=400, detail="User email not found")
        
        # Create checkout session
        result = StripeService.create_checkout_session(
            user_email=user_email,
            tier=stripe_tier,
            success_url=success_url,
            cancel_url=cancel_url,
            user_id=user_id,
        )
        
        return {
            "checkout_url": result['url'],
            "session_id": result['session_id'],
        }
    
    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Error creating checkout session: {e}")
        raise HTTPException(status_code=500, detail="Failed to create checkout session")


@app.post("/stripe/create-query-pack-checkout")
async def create_query_pack_checkout(
    request: dict,
    current_user = Depends(get_current_user)
):
    """
    Create a Stripe Checkout Session for one-time query pack purchase
    
    **Request Body:**
    - pack_type: Pack type (basis = 20 queries for 5€, professional = 50 queries for 10€)
    - success_url: URL to redirect after successful payment
    - cancel_url: URL to redirect if payment is cancelled
    
    **Returns:**
    - checkout_url: URL to redirect user to Stripe Checkout
    - session_id: Stripe Checkout Session ID
    """
    if not STRIPE_ENABLED:
        raise HTTPException(status_code=503, detail="Stripe payment service not configured")
    
    try:
        pack_type = request.get("pack_type", "").lower()
        success_url = request.get("success_url", "")
        cancel_url = request.get("cancel_url", "")
        
        if pack_type not in ["basis", "professional"]:
            raise HTTPException(status_code=400, detail=f"Invalid pack type: {pack_type}")
        
        # Get user info from Firebase token
        user_id = current_user.uid
        user_email = current_user.email
        
        if not user_email:
            raise HTTPException(status_code=400, detail="User email not found")
        
        # Create checkout session for query pack
        result = StripeService.create_query_pack_checkout(
            user_email=user_email,
            pack_type=pack_type,
            success_url=success_url,
            cancel_url=cancel_url,
            user_id=user_id,
        )
        
        return {
            "checkout_url": result['url'],
            "session_id": result['session_id'],
            "queries": result['queries'],
        }
    
    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Error creating query pack checkout: {e}")
        raise HTTPException(status_code=500, detail="Failed to create query pack checkout")


@app.post("/stripe/create-portal-session")
async def create_stripe_portal(
    customer_id: str = Form(...),
    return_url: str = Form(...),
):
    """
    Create a Stripe Customer Portal session
    
    **Parameters:**
    - customer_id: Stripe customer ID
    - return_url: URL to return to after portal session
    
    **Returns:**
    - portal_url: URL to redirect user to Stripe Customer Portal
    """
    if not STRIPE_ENABLED:
        raise HTTPException(status_code=503, detail="Stripe payment service not configured")
    
    try:
        result = StripeService.create_customer_portal_session(
            customer_id=customer_id,
            return_url=return_url,
        )
        
        return {
            "portal_url": result['url'],
        }
    
    except Exception as e:
        logger.error(f"Error creating portal session: {e}")
        raise HTTPException(status_code=500, detail="Failed to create portal session")


@app.get("/stripe/verify-session/{session_id}")
async def verify_checkout_session(
    session_id: str,
    current_user = Depends(get_current_user)
):
    """
    Verify a Stripe Checkout Session status
    
    **Returns:**
    - success: Whether the payment was successful
    - status: Session payment status
    - tier: The purchased tier (if successful)
    """
    if not STRIPE_ENABLED:
        raise HTTPException(status_code=503, detail="Stripe payment service not configured")
    
    try:
        import stripe
        session = stripe.checkout.Session.retrieve(session_id)
        
        # Check if session belongs to this user
        session_user_id = session.metadata.get('user_id')
        if session_user_id != current_user.uid and session_user_id != current_user.email:
            raise HTTPException(status_code=403, detail="Session does not belong to this user")
        
        is_paid = session.payment_status == 'paid'
        tier = session.metadata.get('tier', '').lower()
        
        # Map Stripe tier to frontend tier
        tier_mapping = {
            'tenant': 'basis',
            'pro': 'professional',
            'lawyer': 'lawyer'
        }
        frontend_tier = tier_mapping.get(tier, tier)
        
        return {
            "success": is_paid,
            "status": session.payment_status,
            "tier": frontend_tier if is_paid else None,
        }
    
    except stripe.error.InvalidRequestError:
        raise HTTPException(status_code=404, detail="Session not found")
    except Exception as e:
        logger.error(f"Error verifying checkout session: {e}")
        raise HTTPException(status_code=500, detail="Failed to verify session")


@app.post("/stripe/webhook")
async def stripe_webhook(
    request: Request,
    stripe_signature: str = Header(None, alias="stripe-signature"),
):
    """
    Handle Stripe webhook events
    
    **Events handled:**
    - checkout.session.completed
    - customer.subscription.created
    - customer.subscription.updated
    - customer.subscription.deleted
    - invoice.payment_succeeded
    - invoice.payment_failed
    """
    if not STRIPE_ENABLED:
        raise HTTPException(status_code=503, detail="Stripe payment service not configured")
    
    try:
        # Get raw body
        payload = await request.body()
        
        # Get webhook secret from env
        webhook_secret = settings.stripe_webhook_secret
        if not webhook_secret:
            raise HTTPException(status_code=500, detail="Webhook secret not configured")
        
        # Process webhook
        event_data = StripeService.handle_webhook_event(
            payload=payload,
            signature=stripe_signature,
            webhook_secret=webhook_secret,
        )
        
        if event_data:
            logger.info(f"Processed webhook event: {event_data}")
            
            # Update Firestore based on event type
            await handle_webhook_firestore_update(event_data)
            
            # Send emails based on event type
            if EMAIL_ENABLED:
                await handle_webhook_emails(event_data)
        
        return JSONResponse(content={"status": "success"})
    
    except Exception as e:
        logger.error(f"Webhook error: {e}")
        raise HTTPException(status_code=400, detail=str(e))


async def handle_webhook_firestore_update(event_data: dict):
    """Update Firestore user data based on webhook events."""
    from services.user_service import get_user_service
    
    event_type = event_data.get('event')
    user_id = event_data.get('user_id')
    
    if not user_id:
        logger.warning(f"No user_id in webhook event: {event_type}")
        return
    
    user_service = get_user_service()
    
    try:
        if event_type == 'checkout_completed':
            # Update user subscription after successful checkout
            tier = event_data.get('tier', 'TENANT')
            customer_id = event_data.get('customer_id')
            subscription_id = event_data.get('subscription_id')
            
            success = user_service.update_subscription(
                user_id=user_id,
                tier=tier,
                stripe_customer_id=customer_id,
                stripe_subscription_id=subscription_id,
                subscription_status='active'
            )
            
            if success:
                logger.info(f"✅ Firestore updated for checkout: user={user_id}, tier={tier}")
            else:
                logger.error(f"❌ Failed to update Firestore for checkout: user={user_id}")
                
        elif event_type == 'subscription_created':
            # Subscription created - update tier
            tier = event_data.get('tier', 'TENANT')
            
            user_service.update_subscription(
                user_id=user_id,
                tier=tier,
                subscription_status=event_data.get('status', 'active')
            )
            logger.info(f"✅ Firestore updated for subscription_created: user={user_id}")
            
        elif event_type == 'subscription_updated':
            # Subscription updated (e.g., plan change, renewal)
            status = event_data.get('status', 'active')
            
            # If subscription renewed, reset query count
            if status == 'active':
                user_service.reset_monthly_queries(user_id)
                logger.info(f"✅ Reset monthly queries for user={user_id}")
                
        elif event_type == 'subscription_deleted':
            # Subscription cancelled - downgrade to free
            user_service.cancel_subscription(user_id)
            logger.info(f"✅ Subscription cancelled for user={user_id}")
            
        elif event_type == 'payment_succeeded':
            # Payment succeeded - reset monthly query count
            user_service.reset_monthly_queries(user_id)
            logger.info(f"✅ Payment succeeded, reset queries for user={user_id}")
            
        elif event_type == 'payment_failed':
            # Mark subscription as past_due
            # Don't downgrade yet - Stripe will retry
            logger.warning(f"⚠️ Payment failed for user={user_id}")
            
        elif event_type == 'query_pack_purchased':
            # Query pack purchased - add queries to user's limit
            queries = event_data.get('queries', 0)
            pack_type = event_data.get('pack_type')
            
            success = user_service.add_query_pack(
                user_id=user_id,
                queries=queries
            )
            
            if success:
                logger.info(f"✅ Query pack added: user={user_id}, queries={queries}, pack={pack_type}")
            else:
                logger.error(f"❌ Failed to add query pack: user={user_id}")
            
    except Exception as e:
        logger.error(f"❌ Error updating Firestore for {event_type}: {e}")


async def handle_webhook_emails(event_data: dict):
    """Handle email notifications for webhook events"""
    from services.user_service import get_user_service
    
    event_type = event_data.get('event')
    user_id = event_data.get('user_id')
    
    # Plan names and prices
    PLAN_NAMES = {
        'TENANT': 'Basis',
        'PRO': 'Professional', 
        'LAWYER': 'Lawyer Pro',
    }
    PLAN_PRICES = {
        'TENANT': 19.0,
        'PRO': 39.0,
        'LAWYER': 69.0,
    }
    
    try:
        # Get user details from Firestore
        user_service = get_user_service()
        user_data = user_service.get_user(user_id) if user_id else None
        
        user_email = user_data.get('email', '') if user_data else ''
        user_name = user_data.get('displayName', 'Kunde') if user_data else 'Kunde'
        
        if not user_email:
            logger.warning(f"No email found for user {user_id}")
            return
        
        if event_type == 'checkout_completed':
            # Send order confirmation
            plan_tier = event_data.get('tier', 'TENANT')
            plan_name = PLAN_NAMES.get(plan_tier, plan_tier)
            plan_price = PLAN_PRICES.get(plan_tier, 19.0)
            subscription_id = event_data.get('subscription_id', '')
            
            # Check if B2B customer (has company info)
            company_name = user_data.get('companyName', '') if user_data else ''
            account_type = user_data.get('accountType', '') if user_data else ''
            is_b2b = bool(company_name) or account_type == 'b2b'
            
            if is_b2b and company_name:
                email_service.send_order_confirmation_b2b(
                    user_email=user_email,
                    user_name=user_name,
                    company_name=company_name,
                    plan_name=plan_name,
                    plan_price=plan_price,
                    subscription_id=subscription_id
                )
                logger.info(f"Sent B2B order confirmation to {user_email} ({company_name})")
            else:
                email_service.send_order_confirmation(
                    user_email=user_email,
                    user_name=user_name,
                    plan_name=plan_name,
                    plan_price=plan_price,
                    subscription_id=subscription_id
                )
                logger.info(f"Sent order confirmation to {user_email}")
            
        elif event_type == 'payment_failed':
            # Send payment failed notification
            email_service.send_payment_failed(user_email, user_name)
            
        elif event_type == 'subscription_deleted':
            # Send cancellation confirmation
            from datetime import datetime
            email_service.send_subscription_cancelled(user_email, user_name, datetime.now().strftime('%d.%m.%Y'))
            
    except Exception as e:
        logger.error(f"Error sending webhook email: {e}")


@app.get("/stripe/subscription-status/{subscription_id}")
async def get_subscription_status(subscription_id: str):
    """
    Get subscription status from Stripe
    
    **Parameters:**
    - subscription_id: Stripe subscription ID
    
    **Returns:**
    - status: Subscription status (active, canceled, etc.)
    - current_period_end: End date of current billing period
    - cancel_at_period_end: Whether subscription will cancel at period end
    """
    if not STRIPE_ENABLED:
        raise HTTPException(status_code=503, detail="Stripe payment service not configured")
    
    try:
        result = StripeService.get_subscription_status(subscription_id)
        
        return {
            "status": result['status'],
            "current_period_end": result['current_period_end'].isoformat(),
            "cancel_at_period_end": result['cancel_at_period_end'],
            "customer_id": result['customer_id'],
        }
    
    except Exception as e:
        logger.error(f"Error getting subscription status: {e}")
        raise HTTPException(status_code=500, detail="Failed to get subscription status")


@app.post("/stripe/cancel-subscription/{subscription_id}")
async def cancel_subscription(
    subscription_id: str,
    at_period_end: bool = True,
):
    """
    Cancel a subscription
    
    **Parameters:**
    - subscription_id: Stripe subscription ID
    - at_period_end: If True, cancel at end of billing period; if False, cancel immediately
    
    **Returns:**
    - cancelled: True if successful
    - status: New subscription status
    """
    if not STRIPE_ENABLED:
        raise HTTPException(status_code=503, detail="Stripe payment service not configured")
    
    try:
        result = StripeService.cancel_subscription(
            subscription_id=subscription_id,
            at_period_end=at_period_end,
        )
        
        return result
    
    except Exception as e:
        logger.error(f"Error cancelling subscription: {e}")
        raise HTTPException(status_code=500, detail="Failed to cancel subscription")


# === INVOICE / BILLING ENDPOINTS ===

@app.get("/admin/invoices")
async def get_all_invoices(
    limit: int = 100,
    starting_after: str = None,
):
    """
    Get all invoices from Stripe (Admin only)
    
    **Parameters:**
    - limit: Max number of invoices to return
    - starting_after: Pagination cursor
    
    **Returns:**
    - invoices: List of invoice objects
    """
    if not STRIPE_ENABLED:
        raise HTTPException(status_code=503, detail="Stripe payment service not configured")
    
    try:
        import stripe
        
        params = {
            'limit': min(limit, 100),
            'expand': ['data.customer', 'data.subscription']
        }
        if starting_after:
            params['starting_after'] = starting_after
        
        invoices = stripe.Invoice.list(**params)
        
        result = []
        for inv in invoices.data:
            customer = inv.customer if isinstance(inv.customer, dict) else {'email': 'Unbekannt'}
            result.append({
                'id': inv.id,
                'number': inv.number,
                'customer_email': customer.get('email', 'Unbekannt') if isinstance(customer, dict) else getattr(customer, 'email', 'Unbekannt'),
                'customer_name': customer.get('name', '') if isinstance(customer, dict) else getattr(customer, 'name', ''),
                'customer_id': inv.customer if isinstance(inv.customer, str) else inv.customer.id,
                'amount_due': inv.amount_due / 100,  # Convert cents to EUR
                'amount_paid': inv.amount_paid / 100,
                'currency': inv.currency.upper(),
                'status': inv.status,
                'created': inv.created,
                'due_date': inv.due_date,
                'paid_at': inv.status_transitions.paid_at if inv.status_transitions else None,
                'invoice_pdf': inv.invoice_pdf,
                'hosted_invoice_url': inv.hosted_invoice_url,
                'subscription_id': inv.subscription if isinstance(inv.subscription, str) else (inv.subscription.id if inv.subscription else None),
                'period_start': inv.period_start,
                'period_end': inv.period_end,
            })
        
        return {
            'invoices': result,
            'has_more': invoices.has_more,
            'total_count': len(result)
        }
        
    except Exception as e:
        logger.error(f"Error fetching invoices: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch invoices: {str(e)}")


@app.get("/admin/invoices/{invoice_id}")
async def get_invoice(invoice_id: str):
    """
    Get single invoice details from Stripe
    """
    if not STRIPE_ENABLED:
        raise HTTPException(status_code=503, detail="Stripe payment service not configured")
    
    try:
        import stripe
        
        inv = stripe.Invoice.retrieve(invoice_id, expand=['customer', 'subscription', 'lines'])
        customer = inv.customer
        
        return {
            'id': inv.id,
            'number': inv.number,
            'customer_email': customer.email if customer else 'Unbekannt',
            'customer_name': customer.name if customer else '',
            'amount_due': inv.amount_due / 100,
            'amount_paid': inv.amount_paid / 100,
            'currency': inv.currency.upper(),
            'status': inv.status,
            'created': inv.created,
            'invoice_pdf': inv.invoice_pdf,
            'hosted_invoice_url': inv.hosted_invoice_url,
            'lines': [{
                'description': line.description,
                'amount': line.amount / 100,
                'quantity': line.quantity,
            } for line in inv.lines.data]
        }
        
    except Exception as e:
        logger.error(f"Error fetching invoice: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch invoice: {str(e)}")


@app.post("/admin/invoices/{invoice_id}/void")
async def void_invoice(invoice_id: str):
    """
    Void (stornieren) an invoice - creates credit note
    Only works for open invoices
    """
    if not STRIPE_ENABLED:
        raise HTTPException(status_code=503, detail="Stripe payment service not configured")
    
    try:
        import stripe
        
        inv = stripe.Invoice.retrieve(invoice_id)
        
        if inv.status == 'paid':
            # For paid invoices, create a credit note (Gutschrift)
            credit_note = stripe.CreditNote.create(
                invoice=invoice_id,
                reason='order_change',
                memo='Widerruf - Volle Erstattung',
            )
            
            return {
                'success': True,
                'action': 'credit_note_created',
                'credit_note_id': credit_note.id,
                'amount_credited': credit_note.amount / 100,
                'message': f'Gutschrift erstellt: {credit_note.amount / 100}€'
            }
        
        elif inv.status in ['draft', 'open']:
            # For unpaid invoices, void them
            voided = stripe.Invoice.void_invoice(invoice_id)
            
            return {
                'success': True,
                'action': 'voided',
                'invoice_id': voided.id,
                'message': 'Rechnung storniert'
            }
        
        else:
            return {
                'success': False,
                'message': f'Rechnung mit Status "{inv.status}" kann nicht storniert werden'
            }
        
    except Exception as e:
        logger.error(f"Error voiding invoice: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to void invoice: {str(e)}")


@app.post("/admin/invoices/{invoice_id}/refund")
async def refund_invoice(invoice_id: str, amount: float = None):
    """
    Refund a paid invoice (full or partial)
    
    **Parameters:**
    - invoice_id: Stripe invoice ID
    - amount: Amount to refund (optional, full refund if not specified)
    """
    if not STRIPE_ENABLED:
        raise HTTPException(status_code=503, detail="Stripe payment service not configured")
    
    try:
        import stripe
        
        inv = stripe.Invoice.retrieve(invoice_id)
        
        if inv.status != 'paid':
            raise HTTPException(status_code=400, detail="Nur bezahlte Rechnungen können erstattet werden")
        
        # Get the payment intent
        payment_intent = inv.payment_intent
        
        if not payment_intent:
            raise HTTPException(status_code=400, detail="Keine Zahlungsinformationen gefunden")
        
        # Create refund
        refund_params = {'payment_intent': payment_intent}
        if amount:
            refund_params['amount'] = int(amount * 100)  # Convert EUR to cents
        
        refund = stripe.Refund.create(**refund_params)
        
        return {
            'success': True,
            'refund_id': refund.id,
            'amount_refunded': refund.amount / 100,
            'status': refund.status,
            'message': f'Erstattung von {refund.amount / 100}€ erfolgreich'
        }
        
    except Exception as e:
        logger.error(f"Error refunding invoice: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to refund: {str(e)}")


@app.get("/admin/credit-notes")
async def get_credit_notes(limit: int = 50):
    """
    Get all credit notes (Gutschriften)
    """
    if not STRIPE_ENABLED:
        raise HTTPException(status_code=503, detail="Stripe payment service not configured")
    
    try:
        import stripe
        
        credit_notes = stripe.CreditNote.list(limit=limit, expand=['data.invoice', 'data.customer'])
        
        result = []
        for cn in credit_notes.data:
            result.append({
                'id': cn.id,
                'number': cn.number,
                'invoice_id': cn.invoice.id if cn.invoice else None,
                'invoice_number': cn.invoice.number if cn.invoice else None,
                'customer_email': cn.customer.email if cn.customer else 'Unbekannt',
                'amount': cn.amount / 100,
                'reason': cn.reason,
                'memo': cn.memo,
                'status': cn.status,
                'created': cn.created,
                'pdf': cn.pdf,
            })
        
        return {
            'credit_notes': result,
            'total_count': len(result)
        }
        
    except Exception as e:
        logger.error(f"Error fetching credit notes: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch credit notes: {str(e)}")


@app.get("/user/invoices")
async def get_user_invoices(
    current_user = Depends(get_current_user)
):
    """
    Get invoices for the current user
    """
    if not STRIPE_ENABLED:
        raise HTTPException(status_code=503, detail="Stripe payment service not configured")
    
    try:
        import stripe
        from services.user_service import get_user_service
        
        user_service = get_user_service()
        user_data = user_service.get_user(current_user.uid)
        
        if not user_data:
            return {'invoices': []}
        
        customer_id = user_data.get('stripeCustomerId')
        
        if not customer_id:
            return {'invoices': []}
        
        invoices = stripe.Invoice.list(customer=customer_id, limit=50)
        
        result = []
        for inv in invoices.data:
            result.append({
                'id': inv.id,
                'number': inv.number,
                'amount': inv.amount_paid / 100,
                'currency': inv.currency.upper(),
                'status': inv.status,
                'created': inv.created,
                'paid_at': inv.status_transitions.paid_at if inv.status_transitions else None,
                'invoice_pdf': inv.invoice_pdf,
                'hosted_invoice_url': inv.hosted_invoice_url,
            })
        
        return {'invoices': result}
        
    except Exception as e:
        logger.error(f"Error fetching user invoices: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch invoices: {str(e)}")


# === EMAIL API ===

@app.post("/email/send-welcome")
async def send_welcome_email(
    user_email: str = Form(...),
    user_name: str = Form(None),
):
    """
    Send welcome email to new user
    
    **Parameters:**
    - user_email: User's email address
    - user_name: User's display name (optional)
    
    **Returns:**
    - success: True if email was sent
    """
    if not EMAIL_ENABLED:
        raise HTTPException(status_code=503, detail="Email service not configured")
    
    try:
        success = email_service.send_welcome_email(user_email, user_name)
        
        if success:
            return {"success": True, "message": f"Welcome email sent to {user_email}"}
        else:
            raise HTTPException(status_code=500, detail="Failed to send email")
    
    except Exception as e:
        logger.error(f"Error sending welcome email: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/email/send-delete-request")
async def send_delete_request(
    user_email: str = Form(...),
    user_id: str = Form(...),
    user_name: str = Form(None),
):
    """
    Send account deletion request to admin
    
    **Parameters:**
    - user_email: User's email address
    - user_id: Firebase User ID
    - user_name: User's display name (optional)
    
    **Returns:**
    - success: True if email was sent
    """
    try:
        # Send email to admin
        admin_email = "kontakt@domulex.ai"
        
        success = email_service.send_email(
            to_email=admin_email,
            subject=f"🗑️ Account-Löschanfrage: {user_email}",
            html_content=f"""
            <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                <h2 style="color: #1e3a5f;">Account-Löschanfrage</h2>
                <p>Ein Nutzer hat die Löschung seines Accounts angefordert:</p>
                <table style="width: 100%; border-collapse: collapse; margin: 20px 0;">
                    <tr style="background: #f5f5f5;">
                        <td style="padding: 10px; border: 1px solid #ddd;"><strong>E-Mail:</strong></td>
                        <td style="padding: 10px; border: 1px solid #ddd;">{user_email}</td>
                    </tr>
                    <tr>
                        <td style="padding: 10px; border: 1px solid #ddd;"><strong>Name:</strong></td>
                        <td style="padding: 10px; border: 1px solid #ddd;">{user_name or 'Nicht angegeben'}</td>
                    </tr>
                    <tr style="background: #f5f5f5;">
                        <td style="padding: 10px; border: 1px solid #ddd;"><strong>User ID:</strong></td>
                        <td style="padding: 10px; border: 1px solid #ddd;">{user_id}</td>
                    </tr>
                </table>
                <p style="color: #666;">Bitte löschen Sie den Account über das Admin-Panel.</p>
                <a href="https://domulex.ai/admin" style="display: inline-block; padding: 12px 24px; background: #dc2626; color: white; text-decoration: none; border-radius: 6px; margin-top: 10px;">
                    Zum Admin-Panel
                </a>
            </div>
            """
        )
        
        if success:
            logger.info(f"Delete request sent for user {user_email}")
            return {"success": True, "message": "Löschanfrage wurde gesendet"}
        else:
            raise HTTPException(status_code=500, detail="E-Mail konnte nicht gesendet werden")
    
    except Exception as e:
        logger.error(f"Error sending delete request: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/generate_document", response_model=GeneratedDocument)
async def generate_legal_document(
    request: DocumentRequest,
    user_tier: str = Form(...),
    doc_generator: DocumentGenerator = Depends(get_doc_generator),
    user: dict = Depends(count_ai_query),
):
    """
    Generate legal documents and templates (LAWYER TIER ONLY)
    
    **Tier Restriction:** Lawyer Pro only (49€/month)
    
    **Document Types:**
    - KLAGE: Klageschrift für Gericht
    - MAHNUNG: Zahlungsaufforderung
    - KUENDIGUNG: Kündigungsschreiben (Miete/Vertrag)
    - WIDERSPRUCH: Widerspruchsschreiben
    - MAENGELANZEIGE: Mängelanzeige an Vermieter
    - MIETMINDERUNG: Mietminderungsanzeige
    - SCHRIFTSATZ: Allgemeiner Schriftsatz
    - VOLLMACHT: Vollmacht
    - FRISTSETZUNG: Fristsetzungsschreiben
    - EINSPRUCH: Einspruchsschrift
    
    **Required Fields:**
    - document_type: Type of document to generate
    - case_summary: Summary of the legal case/situation
    - party_plaintiff: Name of plaintiff/sender
    - party_defendant: Name of defendant/recipient
    
    **Optional Fields:**
    - legal_basis: Legal basis (e.g., "§ 543 BGB")
    - deadline_days: Deadline in days
    - amount: Amount in EUR if applicable
    - additional_info: Any additional information
    
    **Example Request:**
    ```json
    {
      "document_type": "MAHNUNG",
      "case_summary": "Mieter hat 3 Monate Miete nicht bezahlt (Juli-September 2024). Mehrfache mündliche Aufforderung blieb erfolglos.",
      "party_plaintiff": "Max Mustermann, Vermieter",
      "party_defendant": "Erika Musterfrau, Mieterin",
      "legal_basis": "§ 535 Abs. 2 BGB",
      "amount": 3600.00,
      "deadline_days": 14
    }
    ```
    
    **Returns:**
    - document_type: Type of generated document
    - title: Document title
    - content: Full document text (ready to print)
    - legal_notes: Important legal notes for the client
    - next_steps: Recommended next steps
    """
    try:
        # Tier check - Document generation only for Lawyer tier
        if user_tier.lower() != 'lawyer':
            raise HTTPException(
                status_code=403,
                detail="Dokumentengenerierung ist nur im Lawyer Pro Tarif (49€) verfügbar. Bitte upgraden Sie Ihren Tarif."
            )
        
        # Generate document
        result = doc_generator.generate_document(request)
        
        logger.info(f"Generated {request.document_type} document for user_tier={user_tier}")
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document generation error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Dokumentengenerierung fehlgeschlagen: {str(e)}"
        )


@app.post("/email/send-order-confirmation")
async def send_order_confirmation_email(
    user_email: str = Form(...),
    user_name: str = Form(...),
    plan_name: str = Form(...),
    plan_price: float = Form(...),
    subscription_id: str = Form(...),
    invoice_url: str = Form(None),
):
    """
    Send order confirmation email after successful payment
    
    **Parameters:**
    - user_email: User's email
    - user_name: User's name
    - plan_name: Name of purchased plan
    - plan_price: Monthly price
    - subscription_id: Stripe subscription ID
    - invoice_url: URL to invoice PDF (optional)
    
    **Returns:**
    - success: True if email was sent
    """
    if not EMAIL_ENABLED:
        raise HTTPException(status_code=503, detail="Email service not configured")
    
    try:
        success = email_service.send_order_confirmation(
            user_email=user_email,
            user_name=user_name,
            plan_name=plan_name,
            plan_price=plan_price,
            subscription_id=subscription_id,
            invoice_url=invoice_url,
        )
        
        if success:
            return {"success": True, "message": f"Order confirmation sent to {user_email}"}
        else:
            raise HTTPException(status_code=500, detail="Failed to send email")
    
    except Exception as e:
        logger.error(f"Error sending order confirmation: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/email/send-admin-notification")
async def send_admin_notification_email(
    user_email: str = Form(...),
    user_name: str = Form(...),
    title: str = Form(...),
    message: str = Form(...),
):
    """Send admin notification email to user"""
    if not EMAIL_ENABLED:
        raise HTTPException(status_code=503, detail="Email service not configured")
    
    try:
        success = email_service.send_admin_notification(user_email, user_name, title, message)
        if success:
            return {"success": True, "message": f"Admin notification sent to {user_email}"}
        else:
            raise HTTPException(status_code=500, detail="Failed to send email")
    except Exception as e:
        logger.error(f"Error sending admin notification: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/email/send-tier-change")
async def send_tier_change_email(
    user_email: str = Form(...),
    user_name: str = Form(...),
    new_tier: str = Form(...),
    queries_limit: str = Form(...),
):
    """Send tier change notification email"""
    if not EMAIL_ENABLED:
        raise HTTPException(status_code=503, detail="Email service not configured")
    
    try:
        success = email_service.send_tier_change(user_email, user_name, new_tier, int(queries_limit))
        if success:
            return {"success": True, "message": f"Tier change notification sent to {user_email}"}
        else:
            raise HTTPException(status_code=500, detail="Failed to send email")
    except Exception as e:
        logger.error(f"Error sending tier change notification: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/email/send-account-deletion-reminder")
async def send_account_deletion_reminder_email(
    user_email: str = Form(...),
    user_name: str = Form(...),
    deletion_date: str = Form(...),
):
    """Send account deletion reminder email (7 days before)"""
    if not EMAIL_ENABLED:
        raise HTTPException(status_code=503, detail="Email service not configured")
    
    try:
        success = email_service.send_account_deletion_reminder(user_email, user_name, deletion_date)
        if success:
            return {"success": True, "message": f"Deletion reminder sent to {user_email}"}
        else:
            raise HTTPException(status_code=500, detail="Failed to send email")
    except Exception as e:
        logger.error(f"Error sending deletion reminder: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# === EMAIL CONFIGURATION API (ADMIN ONLY) ===

@app.get("/admin/email/config")
async def get_email_config():
    """Get current email configuration status (without password)"""
    if not EMAIL_ENABLED:
        return {
            "configured": False,
            "error": "Email service not loaded"
        }
    
    return email_service.get_config_status()


@app.post("/admin/email/config")
async def save_email_config(
    smtp_host: str = Form(...),
    smtp_port: int = Form(...),
    smtp_user: str = Form(...),
    smtp_password: str = Form(...),
    from_email: str = Form(...),
    from_name: str = Form(...),
    use_ssl: bool = Form(True),
):
    """Save email configuration to Firestore"""
    try:
        from google.cloud import firestore
        db = firestore.Client()
        
        # Save config to Firestore
        config_ref = db.collection('settings').document('email_config')
        config_ref.set({
            'smtp_host': smtp_host,
            'smtp_port': smtp_port,
            'smtp_user': smtp_user,
            'smtp_password': smtp_password,
            'from_email': from_email,
            'from_name': from_name,
            'use_ssl': use_ssl,
            'updated_at': firestore.SERVER_TIMESTAMP,
        })
        
        # Reload email service config
        if EMAIL_ENABLED:
            email_service.reload_config()
        
        return {"success": True, "message": "E-Mail-Konfiguration gespeichert"}
    
    except Exception as e:
        logger.error(f"Error saving email config: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/admin/email/test-connection")
async def test_email_connection():
    """Test SMTP connection without sending email"""
    if not EMAIL_ENABLED:
        return {"success": False, "error": "Email service not loaded"}
    
    return email_service.test_connection()


@app.post("/admin/email/send-test")
async def send_test_email(
    to_email: str = Form(...),
):
    """Send a test email to verify configuration"""
    if not EMAIL_ENABLED:
        raise HTTPException(status_code=503, detail="Email service not configured")
    
    result = email_service.send_test_email(to_email)
    
    if result.get("success"):
        return result
    else:
        raise HTTPException(status_code=500, detail=result.get("error", "Failed to send test email"))


@app.post("/admin/email/reload")
async def reload_email_config():
    """Reload email configuration from Firestore"""
    if not EMAIL_ENABLED:
        return {"success": False, "error": "Email service not loaded"}
    
    success = email_service.reload_config()
    return {
        "success": success,
        "config": email_service.get_config_status()
    }


# === EMAIL CLIENT API (Full Send/Receive) ===

@app.get("/admin/email/inbox")
async def get_inbox(folder: str = "INBOX", limit: int = 50):
    """Get emails from inbox"""
    if not EMAIL_CLIENT_ENABLED:
        raise HTTPException(status_code=503, detail="Email client not available")
    
    email_client.reload_config()
    emails = email_client.fetch_emails(folder=folder, limit=limit)
    unread = email_client.get_unread_count(folder)
    
    return {
        "emails": emails,
        "total": len(emails),
        "unread": unread,
        "folder": folder
    }


@app.get("/admin/email/folders")
async def get_email_folders():
    """Get list of email folders"""
    if not EMAIL_CLIENT_ENABLED:
        raise HTTPException(status_code=503, detail="Email client not available")
    
    email_client.reload_config()
    folders = email_client.get_folders()
    return {"folders": folders}


@app.get("/admin/email/message/{email_id}")
async def get_email_message(email_id: str, folder: str = "INBOX"):
    """Get single email with full content"""
    if not EMAIL_CLIENT_ENABLED:
        raise HTTPException(status_code=503, detail="Email client not available")
    
    email_client.reload_config()
    message = email_client.get_email(email_id, folder)
    
    if not message:
        raise HTTPException(status_code=404, detail="Email not found")
    
    return message


@app.post("/admin/email/send")
async def send_email_message(
    to: str = Form(...),
    subject: str = Form(...),
    body: str = Form(...),
    cc: str = Form(None),
    bcc: str = Form(None),
    html: str = Form("true"),  # Whether to use Domulex template
):
    """Send an email with optional Domulex.ai branding template"""
    if not EMAIL_CLIENT_ENABLED:
        raise HTTPException(status_code=503, detail="Email client not available")
    
    email_client.reload_config()
    
    use_template = html.lower() == "true"
    
    if use_template:
        result = email_client.send_with_template(
            to=to,
            subject=subject,
            body_text=body,
            use_template=True
        )
    else:
        result = email_client.send_email(
            to=to,
            subject=subject,
            body_text=body,
            cc=cc,
            bcc=bcc
        )
    
    if result.get("success"):
        return result
    else:
        raise HTTPException(status_code=500, detail=result.get("error", "Failed to send email"))


@app.delete("/admin/email/message/{email_id}")
async def delete_email_message(email_id: str, folder: str = "INBOX"):
    """Delete an email"""
    if not EMAIL_CLIENT_ENABLED:
        raise HTTPException(status_code=503, detail="Email client not available")
    
    email_client.reload_config()
    result = email_client.delete_email(email_id, folder)
    
    if result.get("success"):
        return result
    else:
        raise HTTPException(status_code=500, detail=result.get("error", "Failed to delete email"))


@app.post("/admin/email/message/{email_id}/read")
async def mark_email_as_read(email_id: str, folder: str = "INBOX"):
    """Mark email as read"""
    if not EMAIL_CLIENT_ENABLED:
        raise HTTPException(status_code=503, detail="Email client not available")
    
    email_client.reload_config()
    return email_client.mark_as_read(email_id, folder)


@app.post("/admin/email/message/{email_id}/unread")
async def mark_email_as_unread(email_id: str, folder: str = "INBOX"):
    """Mark email as unread"""
    if not EMAIL_CLIENT_ENABLED:
        raise HTTPException(status_code=503, detail="Email client not available")
    
    email_client.reload_config()
    return email_client.mark_as_unread(email_id, folder)


@app.get("/admin/email/unread-count")
async def get_unread_email_count(folder: str = "INBOX"):
    """Get unread email count"""
    if not EMAIL_CLIENT_ENABLED:
        return {"count": 0}
    
    email_client.reload_config()
    count = email_client.get_unread_count(folder)
    return {"count": count}


class SendTemplateRequest(BaseModel):
    template_name: str
    to: str
    params: dict = {}


@app.post("/admin/email/send-template")
async def send_email_template(request: SendTemplateRequest):
    """Send a predefined email template with branding"""
    if not EMAIL_CLIENT_ENABLED:
        raise HTTPException(status_code=503, detail="Email client not available")
    
    email_client.reload_config()
    
    # Import templates
    from services.email_templates import (
        get_welcome_email, get_order_confirmation_email, get_payment_failed_email,
        get_subscription_cancelled_email, get_admin_notification_email,
        get_tier_change_email, get_deletion_reminder_email,
        get_order_confirmation_b2b_email, get_test_email
    )
    
    template_name = request.template_name
    params = request.params
    
    try:
        # Get template based on name
        if template_name == "welcome":
            email_data = get_welcome_email(params.get("user_name", "Kunde"))
        elif template_name == "order_confirmation":
            email_data = get_order_confirmation_email(
                user_name=params.get("user_name", "Kunde"),
                plan_name=params.get("plan_name", "Premium"),
                plan_price=float(params.get("plan_price", "49.00")),
                subscription_id=params.get("subscription_id", "sub_example123")
            )
        elif template_name == "payment_failed":
            email_data = get_payment_failed_email(params.get("user_name", "Kunde"))
        elif template_name == "subscription_cancelled":
            email_data = get_subscription_cancelled_email(
                user_name=params.get("user_name", "Kunde"),
                end_date=params.get("end_date", "31.01.2026")
            )
        elif template_name == "tier_change":
            email_data = get_tier_change_email(
                user_name=params.get("user_name", "Kunde"),
                new_tier=params.get("new_tier", "Premium"),
                queries_limit=int(params.get("queries_limit", "100"))
            )
        elif template_name == "deletion_reminder":
            email_data = get_deletion_reminder_email(
                user_name=params.get("user_name", "Kunde"),
                deletion_date=params.get("end_date", "11.01.2026")
            )
        elif template_name == "admin_notification":
            email_data = get_admin_notification_email(
                user_name=params.get("user_name", "Kunde"),
                title=params.get("title", "Wichtige Information"),
                message=params.get("message", "Dies ist eine Nachricht.")
            )
        elif template_name == "order_confirmation_b2b":
            email_data = get_order_confirmation_b2b_email(
                user_name=params.get("user_name", "Kunde"),
                company_name=params.get("company_name", "Firma GmbH"),
                plan_name=params.get("plan_name", "Business"),
                plan_price=float(params.get("plan_price", "99.00")),
                subscription_id=params.get("subscription_id", "sub_example123")
            )
        elif template_name == "test":
            email_data = get_test_email()
        else:
            raise HTTPException(status_code=400, detail=f"Unknown template: {template_name}")
        
        # Send email with template HTML
        result = email_client.send_email(
            to=request.to,
            subject=email_data["subject"],
            body_html=email_data["html"],
            body_text=email_data["text"]
        )
        
        if result.get("success"):
            return {"success": True, "message": f"E-Mail-Vorlage '{template_name}' erfolgreich an {request.to} gesendet!"}
        else:
            raise HTTPException(status_code=500, detail=result.get("error", "Failed to send email"))
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to send template email: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# === INGESTION API ===
# Include ingestion router for automated legal content updates
try:
    from ingestion.api import router as ingestion_router
    app.include_router(ingestion_router)
    logger.info("✅ Ingestion API enabled")
except ImportError as e:
    logger.warning(f"⚠️ Ingestion API not available: {e}")


# === CRM & DOCUMENT MANAGEMENT API (LAWYER TIER ONLY) ===

@app.post("/crm/clients", response_model=Client)
async def create_client(
    request: CreateClientRequest,
    user_id: str = Form(...),
    user_tier: str = Form(...),
    crm: CRMService = Depends(get_crm_service),
):
    """
    Create a new client (LAWYER TIER ONLY)
    
    Creates a client with AI-generated summary based on notes.
    """
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="CRM ist nur im Lawyer Pro Tarif verfügbar")
    
    try:
        client = await crm.create_client(lawyer_id=user_id, request=request)
        return client
    except Exception as e:
        logger.error(f"Failed to create client: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/crm/clients", response_model=List[Client])
async def list_clients(
    user_id: str,
    user_tier: str,
    status: Optional[ClientStatus] = None,
    limit: int = 100,
    crm: CRMService = Depends(get_crm_service),
):
    """List all clients for a lawyer."""
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="CRM ist nur im Lawyer Pro Tarif verfügbar")
    
    try:
        clients = await crm.list_clients(lawyer_id=user_id, status=status, limit=limit)
        return clients
    except Exception as e:
        logger.error(f"Failed to list clients: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/crm/clients/{client_id}", response_model=Client)
async def get_client(
    client_id: str,
    user_id: str,
    user_tier: str,
    crm: CRMService = Depends(get_crm_service),
):
    """Get a specific client."""
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="CRM ist nur im Lawyer Pro Tarif verfügbar")
    
    client = await crm.get_client(lawyer_id=user_id, client_id=client_id)
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    return client


@app.put("/crm/clients/{client_id}", response_model=Client)
async def update_client(
    client_id: str,
    request: UpdateClientRequest,
    user_id: str = Form(...),
    user_tier: str = Form(...),
    crm: CRMService = Depends(get_crm_service),
):
    """Update client information."""
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="CRM ist nur im Lawyer Pro Tarif verfügbar")
    
    client = await crm.update_client(lawyer_id=user_id, client_id=client_id, request=request)
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    return client


@app.delete("/crm/clients/{client_id}")
async def delete_client(
    client_id: str,
    user_id: str,
    user_tier: str,
    crm: CRMService = Depends(get_crm_service),
):
    """Archive a client."""
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="CRM ist nur im Lawyer Pro Tarif verfügbar")
    
    success = await crm.delete_client(lawyer_id=user_id, client_id=client_id)
    if not success:
        raise HTTPException(status_code=404, detail="Client not found")
    
    return {"success": True}


# === MANDATE MANAGEMENT ===

@app.post("/crm/mandates", response_model=Mandate)
async def create_mandate(
    request: CreateMandateRequest,
    user_id: str = Form(...),
    user_tier: str = Form(...),
    crm: CRMService = Depends(get_crm_service),
):
    """Create a new mandate/case (LAWYER TIER ONLY)"""
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="CRM ist nur im Lawyer Pro Tarif verfügbar")
    
    try:
        mandate = await crm.create_mandate(lawyer_id=user_id, request=request)
        return mandate
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Failed to create mandate: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/crm/mandates", response_model=List[Mandate])
async def list_mandates(
    user_id: str,
    user_tier: str,
    client_id: Optional[str] = None,
    status: Optional[MandateStatus] = None,
    limit: int = 100,
    crm: CRMService = Depends(get_crm_service),
):
    """List all mandates for a lawyer."""
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="CRM ist nur im Lawyer Pro Tarif verfügbar")
    
    try:
        mandates = await crm.list_mandates(
            lawyer_id=user_id,
            client_id=client_id,
            status=status,
            limit=limit
        )
        return mandates
    except Exception as e:
        logger.error(f"Failed to list mandates: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/crm/mandates/{mandate_id}", response_model=Mandate)
async def get_mandate(
    mandate_id: str,
    user_id: str,
    user_tier: str,
    crm: CRMService = Depends(get_crm_service),
):
    """Get a specific mandate."""
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="CRM ist nur im Lawyer Pro Tarif verfügbar")
    
    mandate = await crm.get_mandate(lawyer_id=user_id, mandate_id=mandate_id)
    if not mandate:
        raise HTTPException(status_code=404, detail="Mandate not found")
    
    return mandate


@app.put("/crm/mandates/{mandate_id}", response_model=Mandate)
async def update_mandate(
    mandate_id: str,
    request: UpdateMandateRequest,
    user_id: str = Form(...),
    user_tier: str = Form(...),
    crm: CRMService = Depends(get_crm_service),
):
    """Update mandate information."""
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="CRM ist nur im Lawyer Pro Tarif verfügbar")
    
    mandate = await crm.update_mandate(lawyer_id=user_id, mandate_id=mandate_id, request=request)
    if not mandate:
        raise HTTPException(status_code=404, detail="Mandate not found")
    
    return mandate


@app.post("/crm/mandates/deadlines", response_model=Mandate)
async def add_deadline(
    request: AddDeadlineRequest,
    user_id: str = Form(...),
    user_tier: str = Form(...),
    crm: CRMService = Depends(get_crm_service),
):
    """Add a deadline to a mandate."""
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="CRM ist nur im Lawyer Pro Tarif verfügbar")
    
    mandate = await crm.add_deadline(lawyer_id=user_id, request=request)
    if not mandate:
        raise HTTPException(status_code=404, detail="Mandate not found")
    
    return mandate


@app.get("/crm/deadlines/upcoming")
async def get_upcoming_deadlines(
    user_id: str,
    user_tier: str,
    days_ahead: int = 30,
    crm: CRMService = Depends(get_crm_service),
):
    """Get all upcoming deadlines across all mandates."""
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="CRM ist nur im Lawyer Pro Tarif verfügbar")
    
    try:
        deadlines = await crm.get_upcoming_deadlines(lawyer_id=user_id, days_ahead=days_ahead)
        return {"deadlines": deadlines, "count": len(deadlines)}
    except Exception as e:
        logger.error(f"Failed to get upcoming deadlines: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# === AI INSIGHTS ===

@app.post("/crm/mandates/{mandate_id}/insights", response_model=MandateInsights)
async def generate_mandate_insights(
    mandate_id: str,
    request: GenerateInsightsRequest,
    user_id: str = Form(...),
    user_tier: str = Form(...),
    crm: CRMService = Depends(get_crm_service),
):
    """
    Generate AI insights for a mandate (LAWYER TIER ONLY)
    
    Generates:
    - Legal strategy recommendations
    - Risk assessment
    - Similar cases
    - Success probability estimation
    """
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="CRM ist nur im Lawyer Pro Tarif verfügbar")
    
    try:
        request.mandate_id = mandate_id  # Ensure consistency
        insights = await crm.generate_mandate_insights(lawyer_id=user_id, request=request)
        return insights
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error(f"Failed to generate insights: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# === DOCUMENT MANAGEMENT ===

@app.get("/crm/documents", response_model=List[Document])
async def list_documents(
    user_id: str,
    user_tier: str,
    client_id: Optional[str] = None,
    mandate_id: Optional[str] = None,
    category: Optional[DocumentCategory] = None,
    limit: int = 100,
    crm: CRMService = Depends(get_crm_service),
):
    """List all documents for a lawyer."""
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="Dokumentenmanagement ist nur im Lawyer Pro Tarif verfügbar")
    
    try:
        from models.crm import Document
        documents = await crm.list_documents(
            lawyer_id=user_id,
            client_id=client_id,
            mandate_id=mandate_id,
            category=category,
            limit=limit
        )
        return documents
    except Exception as e:
        logger.error(f"Failed to list documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/crm/documents/search", response_model=List[DocumentSearchResult])
async def search_documents(
    request: SearchDocumentsRequest,
    user_id: str = Form(...),
    user_tier: str = Form(...),
    crm: CRMService = Depends(get_crm_service),
):
    """
    Search documents with AI (LAWYER TIER ONLY)
    
    Uses AI to semantically search through documents and rank by relevance.
    """
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="Dokumentenmanagement ist nur im Lawyer Pro Tarif verfügbar")
    
    try:
        results = await crm.search_documents_ai(lawyer_id=user_id, request=request)
        return results
    except Exception as e:
        logger.error(f"Failed to search documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/crm/chat", response_model=ChatWithDocumentsResponse)
async def chat_with_documents(
    request: ChatWithDocumentsRequest,
    user_id: str = Form(...),
    user_tier: str = Form(...),
    crm: CRMService = Depends(get_crm_service),
):
    """
    Chat with AI using document/mandate context (LAWYER TIER ONLY)
    
    Ask questions about specific cases, documents, or clients.
    The AI will use the context to provide informed answers.
    """
    if user_tier.lower() != 'lawyer':
        raise HTTPException(status_code=403, detail="CRM ist nur im Lawyer Pro Tarif verfügbar")
    
    try:
        response = await crm.chat_with_documents(lawyer_id=user_id, request=request)
        return response
    except Exception as e:
        logger.error(f"Failed to chat with documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# === TEMPLATE FILLING WITH AI ===

class TemplateFillRequest(BaseModel):
    """Request model for AI template filling."""
    template_content: str
    template_name: str
    instructions: str
    document_text: Optional[str] = None  # Extracted text from uploaded document
    user_tier: Optional[str] = None  # User tier (basis, professional, lawyer)
    use_general_knowledge: Optional[bool] = False  # 🔑 NEU: Allgemeines KI-Wissen statt Datenbank (nur Lawyer)


class TemplateFillResponse(BaseModel):
    """Response model for AI template filling."""
    filled_content: str
    changes_made: List[str]


@app.post("/templates/fill", response_model=TemplateFillResponse)
async def fill_template_with_ai(
    request: TemplateFillRequest,
    user: dict = Depends(count_ai_query)
):
    """
    Fill a legal template using AI based on user instructions or document content.
    
    This endpoint uses Gemini AI to:
    - Extract relevant data from uploaded documents
    - Fill in template placeholders based on instructions
    - Maintain proper legal formatting
    
    **QUELLENNUTZUNG:**
    - **Basis/Professional**: Nur verlässliche Datenbank (Gesetze + Höchstgerichte BGH, BFH, EuGH)
    - **Lawyer (use_general_knowledge=false)**: Datenbank mit allen Quellen
    - **Lawyer (use_general_knowledge=true)**: Allgemeines KI-Wissen (mit Halluzinationswarnung)
    
    Available for all tiers.
    """
    try:
        # Initialize Gemini
        genai.configure(api_key=settings.gemini_api_key)
        model = genai.GenerativeModel('gemini-2.0-flash')
        
        # 🔑 Bestimme Quellennutzung basierend auf Tier und Toggle
        is_lawyer = request.user_tier == 'lawyer'
        use_general_knowledge = is_lawyer and request.use_general_knowledge
        
        if use_general_knowledge:
            # ANWALT-MODUS: Erweiterte KI-Leistung mit öffentlichen Quellen
            system_prompt = """Du bist ein hochspezialisierter Rechtsassistent für professionelle Anwälte, spezialisiert auf deutsches Immobilienrecht.

⚠️ WICHTIG: Du hast Zugriff auf öffentliche Rechtsquellen (BGB, BGH-Urteile, Literatur) und erstellst umfassende, tiefgründige juristische Dokumente.

Deine Aufgabe ist es, juristische Vorlagen professionell auszufüllen und bei Bedarf zu ERWEITERN:

ERWEITERTE FÄHIGKEITEN FÜR ANWÄLTE:
1. 📖 Nutze dein Wissen über deutsche Rechtsprechung (BGH, BFH, OLG) und Literatur (Palandt, MüKo)
2. 📝 Erstelle ausführliche, tiefgründige Formulierungen mit rechtlichen Begründungen
3. ⚖️ Füge relevante Paragraphen, Urteile und Fundstellen hinzu
4. 💡 Wenn der Anwalt erweiterte Formulierungen wünscht, baue das Dokument substanziell aus
5. 🔍 Berücksichtige aktuelle Rechtsprechung und Gesetzeslagen

WICHTIGE REGELN:
1. Ersetze ALLE Platzhalter wie [Name], [Datum], [Adresse] etc. mit passenden Werten
2. Füge bei Bedarf rechtliche Begründungen mit Paragraphen und Rechtsprechung hinzu
3. Wenn "ausführlich", "detailliert" oder "erweitert" gewünscht wird: Baue umfassend aus
4. Nutze juristische Fachsprache auf Anwaltsniveau
5. Behalte professionelle Formatierung bei
6. Antworte NUR mit dem fertigen Dokument, keine Erklärungen drumherum

⚠️ HALLUZINATIONS-WARNUNG: Da du auf allgemeines Wissen zugreifst (nicht nur Datenbank), 
kennzeichne unsichere Informationen mit [Bitte prüfen] und weise auf Verifikationsbedarf hin.

HINWEIS: Du greifst auf öffentlich verfügbare Rechtsquellen zurück, nicht nur auf die interne Datenbank."""
        elif is_lawyer:
            # LAWYER MIT DATENBANK: Erweiterte Leistung mit Datenbankquellen
            system_prompt = """Du bist ein hochspezialisierter Rechtsassistent für professionelle Anwälte, spezialisiert auf deutsches Immobilienrecht.

Du arbeitest mit der GEPRÜFTEN DOMULEX-DATENBANK mit 50.000+ verifizierten Rechtsquellen:
- Gesetze: BGB, WEG, BauGB, GrEStG, EStG, UStG, BetrKV, HeizkostenV, LBOs
- Rechtsprechung: EuGH, BGH, BFH, FG, OLG, LG, AG, VG, OVG
- Kommentare: Palandt, MüKo, Staudinger, Schmidt, Bärmann
- Verwaltung: BMF-Schreiben, EStR, UStAE, GrEStR

WICHTIGE REGELN:
1. Ersetze ALLE Platzhalter wie [Name], [Datum], [Adresse] etc. mit passenden Werten
2. Wenn "ausführlich", "detailliert" oder "erweitert" gewünscht wird: Baue mit Rechtsprechung aus
3. Füge bei Bedarf Paragraphenangaben und Fundstellen hinzu
4. Nutze juristische Fachsprache auf Anwaltsniveau
5. Behalte professionelle Formatierung bei
6. Antworte NUR mit dem fertigen Dokument, keine Erklärungen drumherum

✅ VERLÄSSLICHKEIT: Alle Rechtsquellen stammen aus der geprüften Datenbank."""
        else:
            # BASIS/PROFESSIONAL: Nur verlässliche Höchstgerichte + Gesetze
            system_prompt = """Du bist ein erfahrener Rechtsassistent für deutsches Immobilienrecht.
Deine Aufgabe ist es, juristische Vorlagen mit den bereitgestellten Daten auszufüllen.

⚠️ WICHTIG: Du arbeitest ausschließlich mit VERLÄSSLICHEN, VERBINDLICHEN Quellen:
- Gesetze: BGB, WEG, BauGB, GrEStG, EStG (vollständig)
- Höchstgerichte: EuGH, BGH, BFH (nur höchstrichterliche Rechtsprechung)
- Verwaltungsvorschriften: BMF-Schreiben, EStR

KEINE Instanzgerichte (OLG, LG, AG) - diese können abweichen!
KEINE Literatur/Kommentare - nur Primärquellen!

WICHTIGE REGELN:
1. Ersetze ALLE Platzhalter wie [Name], [Datum], [Adresse] etc. mit passenden Werten
2. Wenn Daten aus einem Dokument extrahiert werden sollen, nutze die relevanten Informationen
3. Wenn keine spezifischen Daten vorliegen, nutze realistische Beispieldaten
4. Behalte die juristische Formatierung und Struktur bei
5. Füge KEINE zusätzlichen Klauseln oder Abschnitte hinzu, es sei denn, es wird explizit angefordert
6. Antworte NUR mit der ausgefüllten Vorlage, keine Erklärungen drumherum
7. Behalte den rechtlichen Hinweis am Ende der Vorlage bei
8. Bei Rechtsangaben: NUR Gesetze und höchstrichterliche Urteile (BGH, BFH, EuGH) zitieren

Bei Mietverträgen, Kaufverträgen etc.:
- Verwende realistische Mietpreise/Kaufpreise für deutsche Verhältnisse
- Setze aktuelle Daten ein
- Nutze plausible Adressen und Namen

✅ VERLÄSSLICHKEIT: Alle Angaben sind verbindlich und höchstrichterlich bestätigt."""

        # Combine context
        context_parts = []
        
        if request.document_text:
            context_parts.append(f"HOCHGELADENES DOKUMENT:\n{request.document_text}")
        
        context_parts.append(f"BENUTZERANWEISUNGEN:\n{request.instructions}")
        context_parts.append(f"VORLAGE ({request.template_name}):\n{request.template_content}")
        
        full_prompt = f"{system_prompt}\n\n" + "\n\n---\n\n".join(context_parts)
        full_prompt += "\n\n---\n\nBitte fülle die Vorlage aus und gib NUR die ausgefüllte Vorlage zurück:"
        
        # Generate filled template
        response = model.generate_content(full_prompt)
        filled_content = response.text.strip()
        
        # Identify changes made
        changes = []
        common_placeholders = [
            ("[Name", "Name ersetzt"),
            ("[Datum", "Datum eingefügt"),
            ("[Adresse", "Adresse eingefügt"),
            ("[Miete", "Mietbetrag eingefügt"),
            ("[Kaution", "Kaution berechnet"),
            ("[Kaufpreis", "Kaufpreis eingefügt"),
            ("[Grundstück", "Grundstücksdaten eingefügt"),
            ("[Notar", "Notardaten eingefügt"),
            ("[Frist", "Frist festgelegt"),
        ]
        
        for placeholder, description in common_placeholders:
            if placeholder in request.template_content and placeholder not in filled_content:
                changes.append(description)
        
        if not changes:
            changes.append("Vorlage wurde mit den angegebenen Daten angepasst")
        
        return TemplateFillResponse(
            filled_content=filled_content,
            changes_made=changes
        )
        
    except Exception as e:
        logger.error(f"Failed to fill template: {e}")
        raise HTTPException(status_code=500, detail=f"Fehler beim Ausfüllen der Vorlage: {str(e)}")


@app.post("/templates/extract-document")
async def extract_document_text(
    file: UploadFile = File(...),
    user: dict = Depends(count_ai_query)
):
    """
    Extract text from uploaded documents (PDF, Word, Images, Text).
    
    Supports:
    - PDF files
    - Word documents (.docx)
    - Images (with OCR via Gemini Vision)
    - Text files
    
    Returns extracted text for template filling.
    """
    try:
        # Read file content
        content = await file.read()
        filename = file.filename.lower() if file.filename else "unknown"
        
        extracted_text = ""
        
        if filename.endswith('.pdf'):
            # Use existing PDF parser (static method)
            try:
                extracted_text = PDFParser.extract_text_from_pdf(content)
            except Exception as pdf_error:
                logger.warning(f"PDF parsing failed, trying Gemini: {pdf_error}")
                # Fallback to Gemini for PDF processing
                genai.configure(api_key=settings.gemini_api_key)
                model = genai.GenerativeModel('gemini-2.0-flash')
                
                import base64
                b64_content = base64.b64encode(content).decode('utf-8')
                
                response = model.generate_content([
                    "Extrahiere den vollständigen Text aus diesem PDF-Dokument. Gib nur den extrahierten Text zurück, keine Erklärungen:",
                    {"mime_type": "application/pdf", "data": b64_content}
                ])
                extracted_text = response.text.strip()
                
        elif filename.endswith('.docx'):
            # Extract text from Word document
            try:
                import io
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(content))
                extracted_text = "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                raise HTTPException(status_code=500, detail="python-docx nicht installiert")
                
        elif filename.endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp')):
            # Use Gemini Vision for OCR
            genai.configure(api_key=settings.gemini_api_key)
            model = genai.GenerativeModel('gemini-2.0-flash')
            
            import base64
            b64_content = base64.b64encode(content).decode('utf-8')
            
            # Determine mime type
            mime_type = "image/png"
            if filename.endswith('.jpg') or filename.endswith('.jpeg'):
                mime_type = "image/jpeg"
            elif filename.endswith('.gif'):
                mime_type = "image/gif"
            elif filename.endswith('.webp'):
                mime_type = "image/webp"
            
            response = model.generate_content([
                "Extrahiere den vollständigen Text aus diesem Bild. Gib nur den extrahierten Text zurück, keine Erklärungen:",
                {"mime_type": mime_type, "data": b64_content}
            ])
            extracted_text = response.text.strip()
            
        elif filename.endswith('.txt'):
            # Plain text file
            extracted_text = content.decode('utf-8', errors='ignore')
            
        else:
            # Try to decode as text
            try:
                extracted_text = content.decode('utf-8', errors='ignore')
            except:
                raise HTTPException(status_code=400, detail="Dateiformat nicht unterstützt")
        
        if not extracted_text:
            raise HTTPException(status_code=400, detail="Kein Text konnte extrahiert werden")
        
        return {
            "success": True,
            "filename": file.filename,
            "text": extracted_text,
            "char_count": len(extracted_text)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to extract document: {e}")
        raise HTTPException(status_code=500, detail=f"Fehler beim Extrahieren: {str(e)}")


# === SEMANTIC DOCUMENT SEARCH ===

from services.document_search import (
    DocumentSearchService,
    GeneratedDocumentSearchRequest,
    GeneratedDocumentSearchResult
)


@app.post("/documents/search", response_model=List[GeneratedDocumentSearchResult])
async def search_generated_documents(request: GeneratedDocumentSearchRequest):
    """
    Semantic search for generated documents using AI.
    
    Searches through user's generated documents and ranks them by relevance
    using Gemini AI. Returns AI-generated summaries and excerpts.
    
    Available for all tiers.
    """
    try:
        # Note: In production, this would fetch documents from Firestore
        # For now, this endpoint expects the frontend to pass document data
        # The actual search will be performed client-side with this endpoint
        # providing the AI ranking capability
        
        logger.info(f"Semantic search request from user {request.user_id}: '{request.query}'")
        
        return []  # Frontend will handle document fetching and call this for ranking
        
    except Exception as e:
        logger.error(f"Failed to search documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/documents/search/rank")
async def rank_documents_by_relevance(
    query: str = Form(...),
    documents: str = Form(...)  # JSON string of documents
):
    """
    Rank documents by relevance to search query using AI.
    
    This endpoint receives a list of documents and uses Gemini AI
    to rank them by semantic relevance to the search query.
    
    Returns ranked results with AI-generated summaries.
    """
    try:
        import json
        
        # Parse documents JSON
        docs_list = json.loads(documents)
        
        # Initialize search service
        search_service = DocumentSearchService(gemini_api_key=settings.gemini_api_key)
        
        # Perform semantic search
        results = await search_service.search_documents(
            documents=docs_list,
            query=query,
            max_results=20
        )
        
        logger.info(f"Ranked {len(results)} documents for query: '{query}'")
        
        return results
        
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid documents JSON")
    except Exception as e:
        logger.error(f"Failed to rank documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# === CRM AI INSIGHTS ===

from services.crm_insights import CRMInsightsService


@app.post("/crm/client/generate-summary")
async def generate_client_summary(
    client_name: str = Form(...),
    case_type: str = Form(None),
    notes: str = Form(None),
    case_notes: str = Form("[]"),  # JSON array of notes
    user: dict = Depends(count_ai_query)
):
    """
    Generate AI summary for a client.
    
    Creates a concise professional summary highlighting key aspects of the case.
    """
    try:
        import json
        
        # Parse case notes
        notes_list = json.loads(case_notes) if case_notes else []
        
        # Initialize insights service
        insights_service = CRMInsightsService(gemini_api_key=settings.gemini_api_key)
        
        # Generate summary
        summary = await insights_service.generate_client_summary(
            client_name=client_name,
            case_type=case_type,
            notes=notes,
            case_notes=notes_list
        )
        
        return {"summary": summary}
        
    except Exception as e:
        logger.error(f"Failed to generate client summary: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/crm/case/analyze")
async def analyze_case(
    client_name: str = Form(...),
    case_type: str = Form(...),
    case_description: str = Form(...),
    case_notes: str = Form("[]"),  # JSON array of notes
    user: dict = Depends(count_ai_query)
):
    """
    Generate comprehensive case analysis with strategic recommendations.
    
    Returns:
        - Analysis of the case
        - Strengths (what favors the client)
        - Weaknesses/Risks
        - Strategic recommendations
        - Next steps to take
    """
    try:
        import json
        
        # Parse case notes
        notes_list = json.loads(case_notes) if case_notes else []
        
        # Initialize insights service
        insights_service = CRMInsightsService(gemini_api_key=settings.gemini_api_key)
        
        # Generate analysis
        analysis = await insights_service.generate_case_analysis(
            client_name=client_name,
            case_type=case_type,
            case_description=case_description,
            case_notes=notes_list
        )
        
        logger.info(f"Generated case analysis for {client_name}")
        
        return analysis
        
    except Exception as e:
        logger.error(f"Failed to analyze case: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/crm/case/risk-assessment")
async def assess_case_risk(
    client_name: str = Form(...),
    case_type: str = Form(...),
    case_description: str = Form(...)
):
    """Generate AI-powered risk assessment for a case."""
    try:
        insights_service = CRMInsightsService(gemini_api_key=settings.gemini_api_key)
        
        risk_assessment = await insights_service.generate_risk_assessment(
            client_name=client_name,
            case_type=case_type,
            case_description=case_description
        )
        
        return {"risk_assessment": risk_assessment}
        
    except Exception as e:
        logger.error(f"Failed to assess risk: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/crm/case/suggest-steps")
async def suggest_next_steps(
    case_type: str = Form(...),
    current_status: str = Form(...),
    case_notes: str = Form("[]")
):
    """Suggest next steps for a case based on current state."""
    try:
        import json
        
        notes_list = json.loads(case_notes) if case_notes else []
        
        insights_service = CRMInsightsService(gemini_api_key=settings.gemini_api_key)
        
        steps = await insights_service.suggest_next_steps(
            case_type=case_type,
            current_status=current_status,
            case_notes=notes_list
        )
        
        return {"next_steps": steps}
        
    except Exception as e:
        logger.error(f"Failed to suggest next steps: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# ADMIN USER MANAGEMENT ENDPOINTS
# ============================================================================

from auth import create_firebase_user, delete_firebase_user, is_admin_email

@app.post("/admin/create-user")
async def admin_create_user(
    email: str = Form(...),
    password: str = Form(...),
    name: str = Form(...),
    tier: str = Form("free"),
    is_admin: bool = Form(False),
    admin_email: str = Form(...),  # E-Mail des angemeldeten Admins zur Verifizierung
    account_type: str = Form("test")  # 'test' = kein Upgrade, 'paying' = normaler Kunde
):
    """
    Erstellt einen neuen Benutzer (Firebase Auth + Firestore).
    Nur für Admins.
    
    account_type:
    - 'test': Test-Kunde ohne Upgrade-Buttons (kostenloser Zugang zum gewählten Tier)
    - 'paying': Normaler Kunde (Stripe Checkout wird separat ausgelöst)
    """
    # Admin-Berechtigung prüfen
    if not is_admin_email(admin_email):
        raise HTTPException(status_code=403, detail="Keine Admin-Berechtigung")
    
    try:
        # Firebase Auth User erstellen
        firebase_user = create_firebase_user(email, password, name)
        uid = firebase_user['uid']
        
        # Firestore Dokument erstellen
        from firebase_admin import firestore
        db = firestore.client()
        
        tier_limits = {
            'free': 3,
            'basis': 25,
            'professional': 250,
            'lawyer': 999999,
        }
        
        # Test-User bekommen isTestUser=True, damit sie keine Upgrade-Buttons sehen
        is_test_user = account_type == 'test'
        
        user_doc = {
            'email': email,
            'name': name,
            'tier': tier,
            'queriesUsed': 0,
            'queriesLimit': tier_limits.get(tier, 3),
            'isAdmin': is_admin,
            'isTestUser': is_test_user,  # Test-Kunden sehen keine Upgrade-Optionen
            'createdAt': firestore.SERVER_TIMESTAMP,
            'updatedAt': firestore.SERVER_TIMESTAMP,
            'lastActivityAt': firestore.SERVER_TIMESTAMP,
        }
        
        db.collection('users').document(uid).set(user_doc)
        
        logger.info(f"✅ Admin erstellt User: {email} (Tier: {tier}, Admin: {is_admin}, TestUser: {is_test_user})")
        
        # Welcome E-Mail senden
        try:
            await send_welcome_email(email, name)
        except Exception as e:
            logger.warning(f"Welcome email failed for {email}: {e}")
        
        return {
            "success": True,
            "uid": uid,
            "email": email,
            "message": f"Benutzer {email} erfolgreich erstellt"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Fehler beim Erstellen des Users: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/admin/create-checkout-for-user")
async def admin_create_checkout_for_user(request: Request):
    """
    Erstellt einen Stripe Checkout-Link für einen vom Admin erstellten User.
    Sendet den Link per E-Mail an den User.
    """
    try:
        data = await request.json()
        user_id = data.get('user_id')
        user_email = data.get('user_email')
        tier = data.get('tier')
        admin_email = data.get('admin_email')
        
        if not is_admin_email(admin_email):
            raise HTTPException(status_code=403, detail="Keine Admin-Berechtigung")
        
        if not all([user_id, user_email, tier]):
            raise HTTPException(status_code=400, detail="user_id, user_email und tier erforderlich")
        
        # Stripe Price IDs aus Settings
        tier_to_price = {
            'basis': settings.stripe_price_tenant,
            'professional': settings.stripe_price_pro,
            'lawyer': settings.stripe_price_lawyer,
        }
        
        price_id = tier_to_price.get(tier)
        if not price_id:
            raise HTTPException(status_code=400, detail=f"Ungültiger Tier: {tier}")
        
        import stripe
        stripe.api_key = settings.stripe_secret_key
        
        # Stripe Customer erstellen oder finden
        customers = stripe.Customer.list(email=user_email, limit=1)
        if customers.data:
            customer = customers.data[0]
        else:
            customer = stripe.Customer.create(
                email=user_email,
                metadata={'firebase_uid': user_id}
            )
        
        # Firestore updaten mit Stripe Customer ID
        from firebase_admin import firestore
        db = firestore.client()
        db.collection('users').document(user_id).update({
            'stripeCustomerId': customer.id
        })
        
        # Checkout Session erstellen
        success_url = f"https://domulex-ai.web.app/dashboard?session_id={{CHECKOUT_SESSION_ID}}"
        cancel_url = "https://domulex-ai.web.app/dashboard"
        
        checkout_session = stripe.checkout.Session.create(
            customer=customer.id,
            payment_method_types=['card'],
            line_items=[{'price': price_id, 'quantity': 1}],
            mode='subscription',
            success_url=success_url,
            cancel_url=cancel_url,
            metadata={
                'firebase_uid': user_id,
                'tier': tier.upper(),
                'created_by_admin': 'true'
            }
        )
        
        # E-Mail mit Checkout-Link senden
        tier_names = {
            'basis': 'Basis (19€/Monat)',
            'professional': 'Professional (39€/Monat)', 
            'lawyer': 'Lawyer Pro (69€/Monat)'
        }
        
        try:
            await send_checkout_invitation_email(
                user_email, 
                user_email.split('@')[0], 
                tier_names.get(tier, tier),
                checkout_session.url
            )
            logger.info(f"📧 Checkout-Link an {user_email} gesendet")
        except Exception as e:
            logger.warning(f"E-Mail an {user_email} fehlgeschlagen: {e}")
        
        logger.info(f"✅ Admin erstellt Checkout für {user_email}: {tier}")
        
        return {
            "success": True,
            "checkout_url": checkout_session.url,
            "message": f"Checkout-Link für {tier} erstellt und per E-Mail gesendet"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Fehler beim Erstellen des Checkout-Links: {e}")
        raise HTTPException(status_code=500, detail=str(e))


async def send_checkout_invitation_email(to_email: str, name: str, tier_name: str, checkout_url: str):
    """Sendet E-Mail mit Checkout-Link an neuen Kunden."""
    subject = f"Ihr Zugang zu domulex.ai - {tier_name}"
    
    html_content = f"""
    <html>
    <body style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px;">
        <div style="background: linear-gradient(135deg, #1e3a5f 0%, #2d4a6f 100%); padding: 30px; border-radius: 10px 10px 0 0;">
            <h1 style="color: #b8860b; margin: 0; font-size: 28px;">domulex.ai</h1>
            <p style="color: white; margin: 10px 0 0 0;">Ihr KI-Assistent für Immobilienrecht</p>
        </div>
        
        <div style="background: #f8f9fa; padding: 30px; border-radius: 0 0 10px 10px;">
            <h2 style="color: #1e3a5f;">Willkommen bei domulex.ai!</h2>
            
            <p>Ihr Account wurde erfolgreich erstellt. Um den <strong>{tier_name}</strong> Tarif zu aktivieren, 
            schließen Sie bitte die Zahlung ab:</p>
            
            <div style="text-align: center; margin: 30px 0;">
                <a href="{checkout_url}" 
                   style="display: inline-block; background: #b8860b; color: white; padding: 15px 40px; 
                          text-decoration: none; border-radius: 8px; font-weight: bold; font-size: 16px;">
                    💳 Jetzt {tier_name} aktivieren
                </a>
            </div>
            
            <p style="color: #666; font-size: 14px;">
                Nach erfolgreicher Zahlung werden alle Premium-Funktionen automatisch freigeschaltet.
            </p>
            
            <hr style="border: none; border-top: 1px solid #ddd; margin: 20px 0;">
            
            <p style="color: #999; font-size: 12px;">
                Diese E-Mail wurde automatisch von domulex.ai versendet.<br>
                Bei Fragen: <a href="mailto:kontakt@domulex.ai">kontakt@domulex.ai</a>
            </p>
        </div>
    </body>
    </html>
    """
    
    await send_email_internal(to_email, subject, html_content)


@app.delete("/admin/delete-user/{user_id}")
async def admin_delete_user(
    user_id: str,
    admin_email: str = Form(...)
):
    """
    Löscht einen Benutzer komplett (Firebase Auth + Firestore).
    Nur für Admins.
    """
    if not is_admin_email(admin_email):
        raise HTTPException(status_code=403, detail="Keine Admin-Berechtigung")
    
    try:
        from firebase_admin import firestore
        db = firestore.client()
        
        # Firestore Dokument löschen
        db.collection('users').document(user_id).delete()
        
        # Notifications löschen
        notifs = db.collection('notifications').where('userId', '==', user_id).stream()
        for notif in notifs:
            notif.reference.delete()
        
        # Firebase Auth User löschen
        delete_firebase_user(user_id)
        
        logger.info(f"✅ Admin löschte User: {user_id}")
        
        return {"success": True, "message": "Benutzer gelöscht"}
        
    except Exception as e:
        logger.error(f"❌ Fehler beim Löschen: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# === DATEN-EXPORT & IMPORT API (PROFESSIONAL/LAWYER TIER) ===

class ExportRequest(BaseModel):
    export_type: str  # 'mandanten', 'objekte', 'all'
    format: str = 'json'  # 'json', 'csv', or 'excel'

class ImportRequest(BaseModel):
    import_type: str  # 'mandanten', 'objekte'
    data: List[dict]
    merge_strategy: str = 'skip'  # 'skip', 'overwrite', 'merge'


@app.post("/data/export")
async def export_user_data(
    request: ExportRequest,
    user: FirebaseUser = Depends(get_current_user),
):
    """
    Export user data (Mandanten/Objekte) as JSON or CSV.
    Requires Professional or Lawyer tier.
    """
    try:
        from firebase_admin import firestore
        db = firestore.client()
        
        # Get user tier from Firestore
        user_doc = db.collection('users').document(user.uid).get()
        user_data = user_doc.to_dict() if user_doc.exists else {}
        user_tier = user_data.get('tier', 'free')
        tier_lower = user_tier.lower()
        
        if tier_lower not in ['professional', 'lawyer']:
            raise HTTPException(status_code=403, detail="Datenexport ist nur im Professional oder Lawyer Pro Tarif verfügbar")
        
        export_data = {
            "export_date": datetime.now().isoformat(),
            "user_id": user.uid,
            "tier": user_tier,
        }
        
        # Export Mandanten (nur Lawyer)
        if request.export_type in ['mandanten', 'all'] and tier_lower == 'lawyer':
            clients_ref = db.collection('users').document(user.uid).collection('clients')
            clients_docs = clients_ref.stream()
            clients = []
            for doc in clients_docs:
                client_data = doc.to_dict()
                client_data['id'] = doc.id
                # Convert timestamps
                for key in ['createdAt', 'updatedAt']:
                    if key in client_data and hasattr(client_data[key], 'isoformat'):
                        client_data[key] = client_data[key].isoformat()
                    elif key in client_data and hasattr(client_data[key], 'seconds'):
                        client_data[key] = datetime.fromtimestamp(client_data[key].seconds).isoformat()
                clients.append(client_data)
            export_data['mandanten'] = clients
            export_data['mandanten_count'] = len(clients)
        
        # Export Objekte (Professional und Lawyer)
        if request.export_type in ['objekte', 'all']:
            objekte_ref = db.collection('users').document(user.uid).collection('objekte')
            objekte_docs = objekte_ref.stream()
            objekte = []
            for doc in objekte_docs:
                objekt_data = doc.to_dict()
                objekt_data['id'] = doc.id
                # Convert timestamps
                for key in ['createdAt', 'updatedAt']:
                    if key in objekt_data and hasattr(objekt_data[key], 'isoformat'):
                        objekt_data[key] = objekt_data[key].isoformat()
                    elif key in objekt_data and hasattr(objekt_data[key], 'seconds'):
                        objekt_data[key] = datetime.fromtimestamp(objekt_data[key].seconds).isoformat()
                objekte.append(objekt_data)
            export_data['objekte'] = objekte
            export_data['objekte_count'] = len(objekte)
        
        # CSV-Format
        if request.format == 'csv':
            import csv
            import io
            
            csv_output = io.StringIO()
            
            if 'mandanten' in export_data and export_data['mandanten']:
                writer = csv.DictWriter(csv_output, fieldnames=['id', 'name', 'email', 'phone', 'address', 'status', 'caseType', 'notes', 'createdAt'])
                writer.writeheader()
                for client in export_data['mandanten']:
                    writer.writerow({
                        'id': client.get('id', ''),
                        'name': client.get('name', ''),
                        'email': client.get('email', ''),
                        'phone': client.get('phone', ''),
                        'address': client.get('address', ''),
                        'status': client.get('status', ''),
                        'caseType': client.get('caseType', ''),
                        'notes': client.get('notes', ''),
                        'createdAt': client.get('createdAt', ''),
                    })
            
            if 'objekte' in export_data and export_data['objekte']:
                if 'mandanten' in export_data:
                    csv_output.write("\n\n")
                writer = csv.DictWriter(csv_output, fieldnames=['id', 'adresse', 'plz', 'ort', 'typ', 'gesamtflaeche', 'gesamteinheiten', 'baujahr', 'heizungstyp', 'createdAt'])
                writer.writeheader()
                for objekt in export_data['objekte']:
                    writer.writerow({
                        'id': objekt.get('id', ''),
                        'adresse': objekt.get('adresse', ''),
                        'plz': objekt.get('plz', ''),
                        'ort': objekt.get('ort', ''),
                        'typ': objekt.get('typ', ''),
                        'gesamtflaeche': objekt.get('gesamtflaeche', ''),
                        'gesamteinheiten': objekt.get('gesamteinheiten', ''),
                        'baujahr': objekt.get('baujahr', ''),
                        'heizungstyp': objekt.get('heizungstyp', ''),
                        'createdAt': objekt.get('createdAt', ''),
                    })
            
            return {
                "success": True,
                "format": "csv",
                "data": csv_output.getvalue()
            }
        
        # Excel-Format
        if request.format == 'excel':
            import io
            import base64
            try:
                import openpyxl
                from openpyxl.styles import Font, PatternFill, Alignment
            except ImportError:
                raise HTTPException(status_code=500, detail="Excel-Export nicht verfügbar. openpyxl fehlt.")
            
            wb = openpyxl.Workbook()
            
            # Header styling
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="1e3a5f", end_color="1e3a5f", fill_type="solid")
            
            # Mandanten Sheet
            if 'mandanten' in export_data and export_data['mandanten']:
                ws = wb.active
                ws.title = "Mandanten"
                headers = ['ID', 'Name', 'E-Mail', 'Telefon', 'Adresse', 'Status', 'Falltyp', 'Notizen', 'Erstellt']
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=header)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center')
                
                for row_idx, client in enumerate(export_data['mandanten'], 2):
                    ws.cell(row=row_idx, column=1, value=client.get('id', ''))
                    ws.cell(row=row_idx, column=2, value=client.get('name', ''))
                    ws.cell(row=row_idx, column=3, value=client.get('email', ''))
                    ws.cell(row=row_idx, column=4, value=client.get('phone', ''))
                    ws.cell(row=row_idx, column=5, value=client.get('address', ''))
                    ws.cell(row=row_idx, column=6, value=client.get('status', ''))
                    ws.cell(row=row_idx, column=7, value=client.get('caseType', ''))
                    ws.cell(row=row_idx, column=8, value=client.get('notes', ''))
                    ws.cell(row=row_idx, column=9, value=client.get('createdAt', ''))
                
                # Auto-width columns
                for col in ws.columns:
                    max_length = max(len(str(cell.value or '')) for cell in col)
                    ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)
            
            # Objekte Sheet
            if 'objekte' in export_data and export_data['objekte']:
                if 'mandanten' in export_data:
                    ws = wb.create_sheet("Objekte")
                else:
                    ws = wb.active
                    ws.title = "Objekte"
                
                headers = ['ID', 'Adresse', 'PLZ', 'Ort', 'Typ', 'Gesamtfläche', 'Einheiten', 'Baujahr', 'Heizungstyp', 'Erstellt']
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=header)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center')
                
                for row_idx, objekt in enumerate(export_data['objekte'], 2):
                    ws.cell(row=row_idx, column=1, value=objekt.get('id', ''))
                    ws.cell(row=row_idx, column=2, value=objekt.get('adresse', ''))
                    ws.cell(row=row_idx, column=3, value=objekt.get('plz', ''))
                    ws.cell(row=row_idx, column=4, value=objekt.get('ort', ''))
                    ws.cell(row=row_idx, column=5, value=objekt.get('typ', ''))
                    ws.cell(row=row_idx, column=6, value=objekt.get('gesamtflaeche', ''))
                    ws.cell(row=row_idx, column=7, value=objekt.get('gesamteinheiten', ''))
                    ws.cell(row=row_idx, column=8, value=objekt.get('baujahr', ''))
                    ws.cell(row=row_idx, column=9, value=objekt.get('heizungstyp', ''))
                    ws.cell(row=row_idx, column=10, value=objekt.get('createdAt', ''))
                
                # Auto-width columns
                for col in ws.columns:
                    max_length = max(len(str(cell.value or '')) for cell in col)
                    ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)
            
            # Save to bytes
            excel_output = io.BytesIO()
            wb.save(excel_output)
            excel_output.seek(0)
            
            # Return as base64
            excel_base64 = base64.b64encode(excel_output.getvalue()).decode('utf-8')
            
            return {
                "success": True,
                "format": "excel",
                "data": excel_base64
            }
        
        logger.info(f"✅ Data export for user {user.uid}: {request.export_type}")
        return {
            "success": True,
            "format": "json",
            "data": export_data
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/data/import")
async def import_user_data(
    request: ImportRequest,
    user: FirebaseUser = Depends(get_current_user),
):
    """
    Import user data (Mandanten/Objekte) from JSON.
    Requires Professional or Lawyer tier.
    """
    try:
        from firebase_admin import firestore
        db = firestore.client()
        
        # Get user tier from Firestore
        user_doc = db.collection('users').document(user.uid).get()
        user_data = user_doc.to_dict() if user_doc.exists else {}
        user_tier = user_data.get('tier', 'free')
        tier_lower = user_tier.lower()
        
        if tier_lower not in ['professional', 'lawyer']:
            raise HTTPException(status_code=403, detail="Datenimport ist nur im Professional oder Lawyer Pro Tarif verfügbar")
        
        if request.import_type == 'mandanten' and tier_lower != 'lawyer':
            raise HTTPException(status_code=403, detail="Mandanten-Import ist nur im Lawyer Pro Tarif verfügbar")
        
        imported = 0
        skipped = 0
        errors = []
        
        if request.import_type == 'mandanten':
            collection_ref = db.collection('users').document(user.uid).collection('clients')
        else:
            collection_ref = db.collection('users').document(user.uid).collection('objekte')
        
        for item in request.data:
            try:
                item_id = item.get('id')
                
                # Check if exists
                if item_id:
                    existing = collection_ref.document(item_id).get()
                    if existing.exists:
                        if request.merge_strategy == 'skip':
                            skipped += 1
                            continue
                        elif request.merge_strategy == 'overwrite':
                            collection_ref.document(item_id).set(item)
                            imported += 1
                            continue
                        elif request.merge_strategy == 'merge':
                            collection_ref.document(item_id).update(item)
                            imported += 1
                            continue
                
                # Add timestamps
                item['createdAt'] = firestore.SERVER_TIMESTAMP
                item['updatedAt'] = firestore.SERVER_TIMESTAMP
                
                # Remove id from data (will be auto-generated)
                if 'id' in item:
                    del item['id']
                
                collection_ref.add(item)
                imported += 1
                
            except Exception as item_error:
                errors.append(f"Fehler bei Eintrag: {str(item_error)}")
        
        logger.info(f"✅ Data import for user {user.uid}: {imported} imported, {skipped} skipped")
        
        return {
            "success": True,
            "imported": imported,
            "skipped": skipped,
            "errors": errors
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Import error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/data/import-excel")
async def import_excel_data(
    file: UploadFile = File(...),
    import_type: str = Form(...),
    user: FirebaseUser = Depends(get_current_user),
):
    """
    Import user data from Excel file (.xlsx).
    Requires Professional or Lawyer tier.
    """
    try:
        import openpyxl
        from firebase_admin import firestore
        import io
        
        db = firestore.client()
        
        # Get user tier from Firestore
        user_doc = db.collection('users').document(user.uid).get()
        user_data = user_doc.to_dict() if user_doc.exists else {}
        user_tier = user_data.get('tier', 'free')
        tier_lower = user_tier.lower()
        
        if tier_lower not in ['professional', 'lawyer']:
            raise HTTPException(status_code=403, detail="Datenimport ist nur im Professional oder Lawyer Pro Tarif verfügbar")
        
        if import_type == 'mandanten' and tier_lower != 'lawyer':
            raise HTTPException(status_code=403, detail="Mandanten-Import ist nur im Lawyer Pro Tarif verfügbar")
        
        # Read Excel file
        content = await file.read()
        wb = openpyxl.load_workbook(io.BytesIO(content))
        ws = wb.active
        
        # Get headers from first row
        headers = [cell.value for cell in ws[1] if cell.value]
        
        # Map German headers to English field names
        header_mapping = {
            'ID': 'id', 'Name': 'name', 'E-Mail': 'email', 'Telefon': 'phone',
            'Adresse': 'address', 'Status': 'status', 'Falltyp': 'caseType',
            'Notizen': 'notes', 'Erstellt': 'createdAt',
            # Objekte
            'PLZ': 'plz', 'Ort': 'ort', 'Typ': 'typ',
            'Gesamtfläche': 'gesamtflaeche', 'Einheiten': 'gesamteinheiten',
            'Baujahr': 'baujahr', 'Heizungstyp': 'heizungstyp'
        }
        
        mapped_headers = [header_mapping.get(h, h.lower() if h else '') for h in headers]
        
        imported = 0
        skipped = 0
        errors = []
        
        if import_type == 'mandanten':
            collection_ref = db.collection('users').document(user.uid).collection('clients')
        else:
            collection_ref = db.collection('users').document(user.uid).collection('objekte')
        
        # Process rows (skip header)
        for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 2):
            try:
                if not any(row):  # Skip empty rows
                    continue
                    
                item = {}
                for col_idx, value in enumerate(row):
                    if col_idx < len(mapped_headers) and mapped_headers[col_idx]:
                        item[mapped_headers[col_idx]] = value if value is not None else ''
                
                item_id = item.get('id')
                
                # Check if exists
                if item_id:
                    existing = collection_ref.document(str(item_id)).get()
                    if existing.exists:
                        skipped += 1
                        continue
                
                # Add timestamps
                item['createdAt'] = firestore.SERVER_TIMESTAMP
                item['updatedAt'] = firestore.SERVER_TIMESTAMP
                
                # Remove id from data
                if 'id' in item:
                    del item['id']
                
                # Convert numeric values
                for key in ['gesamtflaeche', 'gesamteinheiten', 'baujahr']:
                    if key in item and item[key]:
                        try:
                            item[key] = float(item[key]) if '.' in str(item[key]) else int(item[key])
                        except:
                            pass
                
                collection_ref.add(item)
                imported += 1
                
            except Exception as item_error:
                errors.append(f"Zeile {row_idx}: {str(item_error)}")
        
        logger.info(f"✅ Excel import for user {user.uid}: {imported} imported, {skipped} skipped")
        
        return {
            "success": True,
            "imported": imported,
            "skipped": skipped,
            "errors": errors
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Excel import error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
    )
